"""ISO Word (.docx) → Bikeshed or Metanorma AsciiDoc converter.

Reads ISO publication Word files and produces either a specbuild Bikeshed
project or a Metanorma AsciiDoc project suitable for further editing.
"""

from __future__ import annotations

import html
import logging
import re
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

# ---------------------------------------------------------------------------
# Style classification
# ---------------------------------------------------------------------------

_HEADING_LEVEL: dict[str, int] = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    # Clause-1 heading variants found in some ISO templates
    "Scope": 1,
    "Scope Title": 1,
}
_ANNEX_SUB_LEVEL: dict[str, int] = {
    "a2": 2,
    "a3": 3,
    "a4": 4,
    "a5": 5,
}
_BULLET_DEPTH: dict[str, int] = {
    "List Continue 1": 1,
    "List Continue 2": 2,
    "List Continue 3": 3,
    "List Bullet 1": 1,
    "List Bullet 2": 2,
    "List Bullet 3": 3,
    "List Bullet": 1,
}
_NUMBER_DEPTH: dict[str, int] = {
    "List Number 1": 1,
    "List Number 2": 2,
    "List Number 3": 3,
    "List Number": 1,
}
_NOTE_STYLES = {"Note", "Note indent", "NOTE", "NOTE indent", "Notex", "note"}
_CODE_STYLES = {"Code (-)", "Code", "Code Char"}
_EXAMPLE_STYLES = {"Example"}
_SKIP_STYLES = {
    "zzCover",
    "zzCover large",
    "zzCopyright",
    "zzCopyright address",
    "Cover Title_A1",
    "Cover Title_A2",
    "Cover Title_B",
    "Main Title 1",
    "Main Title 2",
}
_FRONT_MATTER_STYLES = {
    "Foreword Title",
    "Intro Title",
    "Foreword Text",
    "Intro Text",
    "Abstract",
    "Foreword",
    "Abstract Title",
}
# Term entry styles map to KIND_TERM
_TERM_STYLES = {"Term entry", "Term(s)"}
# Table body/text paragraph styles — treated as KIND_PARA (documented here
# to avoid any "unknown style" debug paths; the fallback already handles them)
_TABLE_TEXT_STYLES = {"Table Text", "Table Body"}

_MONO_FONTS = {
    "courier new",
    "courier",
    "consolas",
    "lucida console",
    "andale mono",
    "dejavu sans mono",
    "monaco",
}

# Namespace shortcuts used in drawing/image XML
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"


# ---------------------------------------------------------------------------
# Intermediate document model
# ---------------------------------------------------------------------------


@dataclass
class RunSpan:
    text: str
    bold: bool = False
    italic: bool = False
    mono: bool = False
    url: str = ""  # non-empty → hyperlink


@dataclass
class DocBlock:
    kind: str  # see KIND_* constants below
    level: int = 0  # heading level or list depth
    text: str = ""  # plain text (fallback / simple content)
    spans: list[RunSpan] = field(default_factory=list)
    extra: dict = field(default_factory=dict)


KIND_HEADING = "heading"  # regular heading; level 1..5
KIND_ANNEX_H = "annex_heading"  # ANNEX style; extra: marker, title; level=1
KIND_ANNEX_S = "annex_sub"  # a2/a3/a4/a5; level 2..5
KIND_PARA = "para"
KIND_FRONT = "front_matter"  # Foreword/Introduction text
KIND_BULLET = "bullet"  # level 1..3
KIND_NUMBER = "number"  # level 1..3
KIND_NOTE = "note"
KIND_CODE = "code"
KIND_EXAMPLE = "example"
KIND_TERM_NUM = "term_num"
KIND_TERM = "term"
KIND_DEFINITION = "definition"
KIND_REF_NORM = "ref_norm"
KIND_BIB_ENTRY = "bib_entry"
KIND_TABLE = "table"  # extra: rows (list[list[str]]), caption
KIND_FIGURE = "figure"  # extra: image_file, caption
KIND_KEY = "key"  # figure key/legend text
KIND_SKIP = "skip"

# Styles used purely to carry a caption for the *next* block
_CAPTION_STYLES = {"Table title", "Figure title"}


# ---------------------------------------------------------------------------
# Inline run extraction
# ---------------------------------------------------------------------------


def _run_spans(para) -> list[RunSpan]:
    """Extract formatted run spans from a paragraph element, handling hyperlinks."""
    from docx.oxml.ns import qn

    spans: list[RunSpan] = []

    def _process_run(run_elem) -> RunSpan | None:
        text = "".join(n.text or "" for n in run_elem.iter() if n.tag.endswith("}t"))
        text = text.replace("\xa0", "\u00a0")  # preserve non-breaking spaces
        if not text:
            return None
        rpr = run_elem.find(qn("w:rPr"))
        bold = italic = mono = False
        if rpr is not None:
            bold = rpr.find(qn("w:b")) is not None or rpr.find(qn("w:bCs")) is not None
            italic = rpr.find(qn("w:i")) is not None or rpr.find(qn("w:iCs")) is not None
            rf = rpr.find(qn("w:rFonts"))
            if rf is not None:
                fn = (rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi")) or "").lower()
                mono = fn in _MONO_FONTS
            rs = rpr.find(qn("w:rStyle"))
            if rs is not None and "code" in rs.get(qn("w:val"), "").lower():
                mono = True
        return RunSpan(text=text, bold=bold, italic=italic, mono=mono)

    for elem in para._element:
        local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if local == "r":
            s = _process_run(elem)
            if s:
                spans.append(s)
        elif local == "hyperlink":
            r_id = elem.get(qn("r:id"), "")
            url = ""
            try:
                if r_id:
                    url = para.part.rels[r_id].target_ref or ""
            except (KeyError, AttributeError):
                pass
            link_text = "".join(n.text or "" for n in elem.iter() if n.tag.endswith("}t"))
            if link_text:
                spans.append(RunSpan(text=link_text, url=url))
    return spans


def _spans_plain(spans: list[RunSpan]) -> str:
    return "".join(s.text for s in spans)


def _spans_to_bikeshed(spans: list[RunSpan]) -> str:
    parts = []
    for s in spans:
        t = html.escape(s.text)
        if s.url:
            parts.append(f"[{t}]({s.url})" if s.url.startswith("http") else t)
        elif s.mono:
            parts.append(f"`{t}`")
        elif s.bold and s.italic:
            parts.append(f"**_{t}_**")
        elif s.bold:
            parts.append(f"**{t}**")
        elif s.italic:
            parts.append(f"_{t}_")
        else:
            parts.append(t)
    return "".join(parts)


def _spans_to_adoc(spans: list[RunSpan]) -> str:
    parts = []
    for s in spans:
        t = s.text.replace("|", "\\|")
        if s.url:
            parts.append(f"link:{s.url}[{t}]" if s.url.startswith("http") else t)
        elif s.mono:
            parts.append(f"`{t}`")
        elif s.bold and s.italic:
            parts.append(f"**__{t}__**")
        elif s.bold:
            parts.append(f"**{t}**")
        elif s.italic:
            parts.append(f"__{t}__")
        else:
            parts.append(t)
    return "".join(parts)


# ---------------------------------------------------------------------------
# ANNEX heading text parser
# ---------------------------------------------------------------------------


def _parse_annex_text(text: str) -> tuple[str, str]:
    """Returns (marker, title) from ANNEX paragraph text."""
    text = text.strip()
    m = re.search(r"\((normative|informative)\)", text, re.IGNORECASE)
    marker = m.group(1).lower() if m else "normative"
    title = re.sub(r"\s*\((?:normative|informative)\)\s*", "", text, flags=re.IGNORECASE).strip()
    return marker, title


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------


def _extract_images(docx_path: Path, images_dir: Path) -> dict[str, str]:
    """Extract embedded images; return map of relationship-id → filename."""
    images_dir.mkdir(parents=True, exist_ok=True)
    rid_to_file: dict[str, str] = {}
    with zipfile.ZipFile(docx_path) as zf:
        names = zf.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        for mf in media:
            fname = Path(mf).name
            if Path(fname).suffix.lower() in (".emf", ".wmf"):
                logging.warning("EMF/WMF image extracted but may not render in browsers: %s", fname)
            with zf.open(mf) as src:
                (images_dir / fname).write_bytes(src.read())
        rels_path = "word/_rels/document.xml.rels"
        if rels_path in names:
            with zf.open(rels_path) as f:
                tree = ET.parse(f)
            ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
            media_names = {Path(m).name for m in media}
            for rel in tree.findall("r:Relationship", ns):
                rid = rel.get("Id", "")
                tgt = Path(rel.get("Target", "")).name
                if tgt in media_names:
                    rid_to_file[rid] = tgt
    return rid_to_file


def _find_image_rid(para_elem) -> str | None:
    """Return the r:embed relationship ID for a drawing in a paragraph."""
    for elem in para_elem.iter():
        if elem.tag.endswith("}blip"):
            embed = elem.get(f"{{{_NS_R}}}embed")
            if embed:
                return embed
    return None


# ---------------------------------------------------------------------------
# Document parsing → list[DocBlock]
# ---------------------------------------------------------------------------


def parse_docx(
    docx_path: Path,
    images_dir: Path | None = None,
    style_overrides: dict[str, str] | None = None,
) -> tuple[list[DocBlock], dict]:
    """Parse a DOCX file into a flat list of DocBlocks plus metadata dict."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for ISO DOCX conversion. "
            "Install it with: pip install python-docx"
        ) from None
    doc = Document(str(docx_path))
    _heading_level = _HEADING_LEVEL
    _note_styles = _NOTE_STYLES
    _bullet_depth = _BULLET_DEPTH
    _number_depth = _NUMBER_DEPTH
    _skip_styles = _SKIP_STYLES
    if style_overrides:
        _skip_s: set[str] = set()
        _note_s: set[str] = set()
        _head_l: dict[str, int] = {}
        _bull_d: dict[str, int] = {}
        _num_d: dict[str, int] = {}
        for k, v in style_overrides.items():
            parts = v.split()
            kind = parts[0] if parts else ""
            level_tok = parts[-1] if len(parts) > 1 else ""
            if kind == KIND_SKIP:
                _skip_s.add(k)
            elif kind == KIND_NOTE:
                _note_s.add(k)
            elif kind == KIND_HEADING and level_tok.isdigit():
                _head_l[k] = int(level_tok)
            elif kind == KIND_BULLET and level_tok.isdigit():
                _bull_d[k] = int(level_tok)
            elif kind == KIND_NUMBER and level_tok.isdigit():
                _num_d[k] = int(level_tok)
        _skip_styles = _SKIP_STYLES | _skip_s
        _note_styles = _NOTE_STYLES | _note_s
        _heading_level = _HEADING_LEVEL | _head_l
        _bullet_depth = _BULLET_DEPTH | _bull_d
        _number_depth = _NUMBER_DEPTH | _num_d
    rid_to_file: dict[str, str] = {}
    if images_dir is not None:
        rid_to_file = _extract_images(docx_path, images_dir)

    meta: dict = {}
    blocks: list[DocBlock] = []
    pending_caption: str | None = None
    pending_fig_caption: str | None = None
    fig_counter = [0]

    # Collect cover metadata from early paragraphs
    for p in doc.paragraphs[:20]:
        sn = p.style.name
        t = p.text.strip().replace("\xa0", " ")
        if sn == "zzCover large" and t:
            meta["doc_ref"] = t
        elif sn in ("Cover Title_A1", "Cover Title_A2") and t:
            meta.setdefault("title_parts", []).append(t)
        elif sn == "zzCover" and re.match(r"\d{4}", t):
            meta["date"] = t

    if "title_parts" in meta:
        meta["title"] = " ".join(meta["title_parts"]).replace("\n", " ").replace("  ", " ")
        # Extract doc number and part number from doc_ref like "ISO/IEC 23000-19:2024(en)"
        m = re.search(r"(\d+)-(\d+):\d+", meta.get("doc_ref", ""))
        if m:
            meta["docnumber"] = m.group(1)
            meta["partnumber"] = m.group(2)

    for elem in doc.element.body:
        local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        if local == "p":
            from docx.text.paragraph import Paragraph as _Para

            para = _Para(elem, doc)
            sn = para.style.name
            spans = _run_spans(para)
            plain = _spans_plain(spans).strip()

            if sn in _skip_styles or not plain:
                continue

            if sn == "Table title":
                pending_caption = plain
                continue
            if sn == "Figure title":
                # applies to the most recent figure block
                if blocks and blocks[-1].kind == KIND_FIGURE:
                    blocks[-1].extra["caption"] = plain
                    blocks[-1].text = plain
                else:
                    if pending_fig_caption is not None:
                        logging.warning(
                            "parse_docx: consecutive Figure title paragraphs with no intervening "
                            "graphic — overwriting %r with %r",
                            pending_fig_caption,
                            plain,
                        )
                    pending_fig_caption = plain
                continue

            # Front matter
            if sn in _FRONT_MATTER_STYLES:
                blocks.append(DocBlock(KIND_FRONT, text=plain, spans=spans))
                continue

            # Regular headings
            if sn in _heading_level:
                blocks.append(
                    DocBlock(KIND_HEADING, level=_heading_level[sn], text=plain, spans=spans)
                )
                continue

            # ANNEX main heading
            if sn == "ANNEX":
                marker, title = _parse_annex_text(plain)
                blocks.append(DocBlock(KIND_ANNEX_H, level=1, text=title, extra={"marker": marker}))
                continue

            # ANNEX sub-headings
            if sn in _ANNEX_SUB_LEVEL:
                blocks.append(
                    DocBlock(KIND_ANNEX_S, level=_ANNEX_SUB_LEVEL[sn], text=plain, spans=spans)
                )
                continue

            # Lists
            if sn in _BULLET_DEPTH:
                blocks.append(
                    DocBlock(KIND_BULLET, level=_BULLET_DEPTH[sn], text=plain, spans=spans)
                )
                continue
            if sn in _NUMBER_DEPTH:
                blocks.append(
                    DocBlock(KIND_NUMBER, level=_NUMBER_DEPTH[sn], text=plain, spans=spans)
                )
                continue

            # Notes
            if sn in _NOTE_STYLES:
                # Strip "NOTE", "NOTE 1", "NOTE 1 to entry:" prefixes
                note_text = re.sub(r"^NOTE\s*\d*\s*(?:to\s+entry\s*)?[:\t]?\s*", "", plain).strip()
                blocks.append(DocBlock(KIND_NOTE, text=note_text or plain, spans=spans))
                continue

            # Code blocks
            if sn in _CODE_STYLES:
                blocks.append(DocBlock(KIND_CODE, text=plain))
                continue

            # Examples
            if sn in _EXAMPLE_STYLES:
                example_text = re.sub(r"^EXAMPLE\s*\d*\s*[–—:\t]?\s*", "", plain).strip()
                blocks.append(DocBlock(KIND_EXAMPLE, text=example_text or plain, spans=spans))
                continue

            # Terms and definitions
            if sn == "TermNum":
                blocks.append(DocBlock(KIND_TERM_NUM, text=plain))
                continue
            if sn in _TERM_STYLES:
                blocks.append(DocBlock(KIND_TERM, text=plain, spans=spans))
                continue
            if sn == "Definition":
                blocks.append(DocBlock(KIND_DEFINITION, text=plain, spans=spans))
                continue

            # Normative references and bibliography
            if sn == "RefNorm":
                blocks.append(DocBlock(KIND_REF_NORM, text=plain, spans=spans))
                continue
            if sn == "Biblio Entry":
                blocks.append(DocBlock(KIND_BIB_ENTRY, text=plain, spans=spans))
                continue

            # Figure graphic (drawing container — no text)
            if sn == "Figure Graphic":
                rid = _find_image_rid(elem) or ""
                img_file = rid_to_file.get(rid, "") if rid_to_file else ""
                if not rid or (rid_to_file and not img_file):
                    cap_hint = pending_fig_caption or f"Figure {fig_counter[0] + 1}"
                    logging.warning(
                        "Figure has no extractable image (possibly EMF/WMF): %s", cap_hint
                    )
                fig_counter[0] += 1
                cap = pending_fig_caption or f"Figure {fig_counter[0]}"
                pending_fig_caption = None
                blocks.append(
                    DocBlock(
                        KIND_FIGURE,
                        text=cap,
                        extra={
                            "image_file": img_file,
                            "caption": cap,
                            "fig_id": f"fig-{fig_counter[0]}",
                            "image_rid": rid,
                        },
                    )
                )
                continue

            # Key/legend
            if sn in ("Key Title", "Key text", "Key"):
                blocks.append(DocBlock(KIND_KEY, text=plain, spans=spans))
                continue

            # Body text (default)
            blocks.append(DocBlock(KIND_PARA, text=plain, spans=spans))

        elif local == "tbl":
            from docx.table import Table as _Table

            table = _Table(elem, doc)
            rows: list[list[str]] = []
            first_row_bold: list[bool] = []
            has_header = False
            for row_idx, row in enumerate(table.rows):
                cells = []
                seen_tcs: set[int] = set()
                row_bold_flags: list[bool] = []
                for cell in row.cells:
                    if id(cell._tc) in seen_tcs:
                        continue
                    seen_tcs.add(id(cell._tc))
                    cell_text = " ".join(
                        cp.text.strip() for cp in cell.paragraphs if cp.text.strip()
                    )
                    cells.append(cell_text.replace("\xa0", " "))
                    # Check if all runs in the cell are bold (header detection)
                    cell_bold = (
                        all(
                            run.bold
                            for cp in cell.paragraphs
                            for run in cp.runs
                            if run.text.strip()
                        )
                        if any(run.text.strip() for cp in cell.paragraphs for run in cp.runs)
                        else False
                    )
                    row_bold_flags.append(cell_bold)
                rows.append(cells)
                if row_idx == 0:
                    first_row_bold = row_bold_flags
            # Detect header: first row all-bold, or table has a header row
            if first_row_bold and all(first_row_bold):
                has_header = True
            elif table.rows and table.rows[0].cells:
                try:
                    tr_pr = table.rows[0]._tr.trPr
                    if tr_pr is not None:
                        from docx.oxml.ns import qn as _qn

                        if tr_pr.find(_qn("w:tblHeader")) is not None:
                            has_header = True
                except (AttributeError, TypeError):
                    pass
            if rows:
                blocks.append(
                    DocBlock(
                        KIND_TABLE,
                        text=pending_caption or "",
                        extra={"rows": rows, "caption": pending_caption, "has_header": has_header},
                    )
                )
            pending_caption = None

    if pending_fig_caption:
        blocks.append(DocBlock(KIND_PARA, text=f"[Figure: {pending_fig_caption}]"))

    if pending_caption is not None:
        logging.warning(
            "parse_docx: document ended with an unused table caption: %r", pending_caption
        )
        blocks.append(DocBlock(KIND_PARA, text=f"[Table: {pending_caption}]"))

    return blocks, meta


# ---------------------------------------------------------------------------
# Slug utilities (shared with other converters)
# ---------------------------------------------------------------------------


def _slugify(text: str) -> str:
    s = text.lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_]+", "-", s)
    s = s.strip("-")
    return s or "section"


# ---------------------------------------------------------------------------
# Split blocks into sections (by top-level headings / annexes)
# ---------------------------------------------------------------------------


@dataclass
class Section:
    slug: str
    title: str
    prefix: str  # "00", "01", "a01", etc.
    blocks: list[DocBlock] = field(default_factory=list)


def _split_sections(blocks: list[DocBlock]) -> list[Section]:
    """Split flat block list into sections at Heading 1 and ANNEX boundaries."""
    sections: list[Section] = []
    current: Section | None = None

    heading_counter = 0
    annex_counter = 0
    annex_letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    front_matter_done = False

    def _push():
        if current is not None:
            sections.append(current)

    for b in blocks:
        if b.kind == KIND_FRONT and not front_matter_done:
            if current is None or current.prefix == "00":
                if current is None:
                    current = Section("foreword", "Foreword", "00")
                current.blocks.append(b)
            else:
                current.blocks.append(b)
            continue

        if b.kind == KIND_HEADING and b.level == 1:
            front_matter_done = True
            _push()
            heading_counter += 1
            slug = _slugify(b.text)
            prefix = f"{heading_counter:02d}"
            current = Section(slug, b.text, prefix)
            continue

        if b.kind == KIND_ANNEX_H:
            front_matter_done = True
            _push()
            annex_counter += 1
            letter = (
                annex_letters[annex_counter - 1] if annex_counter <= 26 else f"X{annex_counter}"
            )
            slug = f"annex-{letter.lower()}"
            prefix = f"a{annex_counter:02d}"
            current = Section(slug, b.text, prefix, blocks=[b])
            continue

        if current is None:
            current = Section("front-matter", "Front Matter", "00")
        current.blocks.append(b)

    _push()
    return sections


# ---------------------------------------------------------------------------
# Bikeshed writer
# ---------------------------------------------------------------------------


def _render_bikeshed_block(b: DocBlock, in_annex: bool = False) -> list[str]:
    lines: list[str] = []

    def _inline(b_: DocBlock) -> str:
        return _spans_to_bikeshed(b_.spans) if b_.spans else b_.text

    if b.kind == KIND_HEADING:
        depth = b.level
        hashes = "#" * depth
        slug = _slugify(b.text)
        lines += ["", f"{hashes} {b.text} {hashes} {{#{slug}}}", ""]

    elif b.kind == KIND_ANNEX_S:
        depth = b.level
        hashes = "#" * depth
        slug = _slugify(b.text)
        lines += ["", f"{hashes} {b.text} {hashes} {{#{slug}}}", ""]

    elif b.kind in (KIND_PARA, KIND_FRONT):
        lines += ["", _inline(b), ""]

    elif b.kind == KIND_BULLET:
        prefix = "  " * (b.level - 1) + "- "
        lines.append(f"{prefix}{_inline(b)}")

    elif b.kind == KIND_NUMBER:
        prefix = "  " * (b.level - 1) + "1. "
        lines.append(f"{prefix}{_inline(b)}")

    elif b.kind == KIND_NOTE:
        lines += ["", f'<div class="note">NOTE — {_inline(b)}</div>', ""]

    elif b.kind == KIND_CODE:
        lines += ["", "<pre>", b.text, "</pre>", ""]

    elif b.kind == KIND_EXAMPLE:
        lines += ["", '<div class="example">', "", _inline(b), "", "</div>", ""]

    elif b.kind == KIND_TERM_NUM:
        lines.append(f"<!-- term {b.text} -->")

    elif b.kind == KIND_TERM:
        lines += ["", f": <dfn>{_inline(b)}</dfn>"]

    elif b.kind == KIND_DEFINITION:
        lines += [f":: {_inline(b)}", ""]

    elif b.kind in (KIND_REF_NORM, KIND_BIB_ENTRY):
        inline_text = _spans_to_bikeshed(b.spans) if b.spans else html.escape(b.text)
        lines.append(f'<li class="reference">{inline_text}</li>')

    elif b.kind == KIND_TABLE:
        cap = b.extra.get("caption", "")
        rows = b.extra.get("rows", [])
        has_header = b.extra.get("has_header", False)
        lines.append("")
        lines.append('<table class="data">')
        if cap:
            lines.append(f"<caption>{cap}</caption>")
        if rows:
            if has_header:
                lines.append("<thead><tr>")
                for cell in rows[0]:
                    lines.append(f"<th>{html.escape(cell)}</th>")
                lines.append("</tr></thead>")
                lines.append("<tbody>")
                for row in rows[1:]:
                    lines.append("<tr>")
                    for cell in row:
                        lines.append(f"<td>{html.escape(cell)}</td>")
                    lines.append("</tr>")
                lines.append("</tbody>")
            else:
                lines.append("<tbody>")
                for row in rows:
                    lines.append("<tr>")
                    for cell in row:
                        lines.append(f"<td>{html.escape(cell)}</td>")
                    lines.append("</tr>")
                lines.append("</tbody>")
        lines += ["</table>", ""]

    elif b.kind == KIND_FIGURE:
        img = b.extra.get("image_file", "")
        cap = b.extra.get("caption", b.text)
        fig_id = b.extra.get("fig_id", "")
        id_attr = f' id="{fig_id}"' if fig_id else ""
        lines += ["", f"<figure{id_attr}>"]
        if img:
            lines.append(f'<img src="images/{img}" alt="{cap}">')
        else:
            lines.append(f"<!-- figure: {cap} -->")
        if cap:
            lines.append(f"<figcaption>{cap}</figcaption>")
        lines += ["</figure>", ""]

    elif b.kind == KIND_KEY:
        lines.append(f"<!-- key: {b.text} -->")

    return lines


def _wrap_reference_lists(lines: list[str]) -> list[str]:
    """Group consecutive ``<li class="reference">`` lines inside a ``<ul class="references">`` wrapper.

    This ensures Bikeshed renders a single contiguous list even when notes or
    other blocks appear between individual reference entries.
    """
    result: list[str] = []
    in_ref = False
    for line in lines:
        is_ref_li = line.startswith('<li class="reference">')
        if is_ref_li and not in_ref:
            result.append('<ul class="references">')
            in_ref = True
        elif not is_ref_li and in_ref:
            result.append("</ul>")
            result.append("")
            in_ref = False
        result.append(line)
    if in_ref:
        result.append("</ul>")
        result.append("")
    return result


def write_bikeshed(
    blocks: list[DocBlock], meta: dict, output_dir: Path, docx_path: Path, overwrite: bool = False
) -> dict:
    """Write a Bikeshed project to output_dir."""
    bs_dir = output_dir / "bikeshed"
    if bs_dir.exists() and not overwrite:
        raise FileExistsError(f"Output directory already exists: {bs_dir}")
    bs_dir.mkdir(parents=True, exist_ok=True)

    # Extract images alongside
    images_dir = bs_dir / "images"
    images_dir.mkdir(exist_ok=True)
    rid_to_file = _extract_images(docx_path, images_dir)
    for b in blocks:
        if b.kind == KIND_FIGURE and not b.extra.get("image_file"):
            rid = b.extra.get("image_rid", "")
            if rid:
                b.extra["image_file"] = rid_to_file.get(rid, "")

    sections = _split_sections(blocks)
    manifest_lines: list[str] = []
    written: list[str] = []

    title = meta.get("title", "Untitled Specification")
    date = meta.get("date", "")

    annex_letter = iter("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
    has_prefix_00 = any(s.prefix == "00" for s in sections)

    for sec in sections:
        # Build filename
        slug_part = re.sub(r"[^a-z0-9]+", "_", sec.title.lower())[:40].strip("_")
        fname = f"{sec.prefix}_{slug_part}.bs"
        manifest_lines.append(fname)
        written.append(fname)

        out_lines: list[str] = []

        # First file gets Bikeshed metadata
        if sec.prefix == "00" or (not has_prefix_00 and sec == sections[0]):
            out_lines += [
                "<!-- Bikeshed metadata -->",
                "<pre class='metadata'>",
                f"Title: {title}",
                "Status: WD",
            ]
            if date:
                out_lines.append(f"Date: {date}")
            out_lines += [
                "Editor: See document",
                "Abstract: " + title,
                "Markup Shorthands: markdown yes",
                "</pre>",
                "",
            ]

        # Section heading
        if sec.prefix != "00" and sec.blocks and sec.blocks[0].kind != KIND_ANNEX_H:
            slug = _slugify(sec.title)
            out_lines += [f"## {sec.title} ## {{#{slug}}}", ""]

        elif (ann_b := next((b for b in sec.blocks if b.kind == KIND_ANNEX_H), None)) is not None:
            letter = next(annex_letter, "X")
            marker = ann_b.extra.get("marker", "normative")
            slug = _slugify(ann_b.text)
            out_lines += [
                f"## Annex {letter} ({marker}) — {ann_b.text} ## {{#{slug}}}",
                "",
            ]
            # Render remaining blocks (skip the annex heading itself)
            for b in sec.blocks:
                if b.kind != KIND_ANNEX_H:
                    out_lines += _render_bikeshed_block(b, in_annex=True)
            out_lines = _wrap_reference_lists(out_lines)
            (bs_dir / fname).write_text("\n".join(out_lines), encoding="utf-8")
            continue

        for b in sec.blocks:
            out_lines += _render_bikeshed_block(b)

        out_lines = _wrap_reference_lists(out_lines)
        (bs_dir / fname).write_text("\n".join(out_lines), encoding="utf-8")

    # manifest.txt
    (bs_dir / "manifest.txt").write_text("\n".join(manifest_lines) + "\n", encoding="utf-8")

    # specbuild.toml
    spec_name = re.sub(
        r"[^A-Za-z0-9_-]", "_", (meta.get("docnumber", title.split()[-1] if title else "Spec"))
    )[:24]
    toml = f"""[standards]
flavor = "iso"
doc_number = "{meta.get("docnumber", "")}-{meta.get("partnumber", "")}"

[spec]
spec_name = "{spec_name}"
spec_full_name = "{title}"
bikeshed_dir = "bikeshed"
"""
    (output_dir / "specbuild.toml").write_text(toml, encoding="utf-8")

    return {"sections": written, "warnings": [], "output_dir": bs_dir}


# ---------------------------------------------------------------------------
# Metanorma AsciiDoc writer
# ---------------------------------------------------------------------------


def _render_adoc_block(b: DocBlock, in_annex: bool = False) -> list[str]:
    lines: list[str] = []

    def _inline(b_: DocBlock) -> str:
        return _spans_to_adoc(b_.spans) if b_.spans else b_.text

    if b.kind == KIND_HEADING:
        prefix = "=" * (b.level + 1)
        slug = _slugify(b.text)
        lines += ["", f"[[{slug}]]", f"{prefix} {b.text}", ""]

    elif b.kind == KIND_ANNEX_S:
        prefix = "=" * (b.level + 1)
        slug = _slugify(b.text)
        lines += ["", f"[[{slug}]]", f"{prefix} {b.text}", ""]

    elif b.kind in (KIND_PARA, KIND_FRONT):
        lines += ["", _inline(b), ""]

    elif b.kind == KIND_BULLET:
        prefix = "*" * b.level + " "
        lines.append(f"{prefix}{_inline(b)}")

    elif b.kind == KIND_NUMBER:
        prefix = "." * b.level + " "
        lines.append(f"{prefix}{_inline(b)}")

    elif b.kind == KIND_NOTE:
        lines += ["", "NOTE: " + _inline(b), ""]

    elif b.kind == KIND_CODE:
        lines += ["", "[source]", "----", b.text, "----", ""]

    elif b.kind == KIND_EXAMPLE:
        lines += ["", "[example]", "====", _inline(b), "====", ""]

    elif b.kind == KIND_TERM_NUM:
        lines.append(f"// term number: {b.text}")

    elif b.kind == KIND_TERM:
        level = b.level if b.level > 0 else 3
        prefix = "=" * (level + 1)
        lines += ["", f"{prefix} {_inline(b)}", ""]

    elif b.kind == KIND_DEFINITION:
        lines += [_inline(b), ""]

    elif b.kind == KIND_BIB_ENTRY:
        text = _inline(b)
        m = re.match(r"^\[(\d+)\]\s*(.*)", text)
        if m:
            num, rest = m.group(1), m.group(2).strip()
            lines.append(f"* [[[bib-{num},{rest.split(',')[0].strip().replace(':', '-')}]]] {rest}")
        else:
            lines.append(f"* {text}")
    elif b.kind == KIND_REF_NORM:
        lines.append(f"* {_inline(b)}")

    elif b.kind == KIND_TABLE:
        cap = b.extra.get("caption", "")
        rows = b.extra.get("rows", [])
        has_header = b.extra.get("has_header", False)
        lines.append("")
        if cap:
            lines.append(f".{cap}")
        ncols = max((len(r) for r in rows), default=1)
        col_attr = f'cols="{",".join(["1"] * ncols)}"'
        if has_header:
            lines.append(f'[{col_attr},options="header"]')
        else:
            lines.append(f"[{col_attr}]")
        lines.append("|===")
        if rows:
            if has_header:
                # Header row — plain `| cell` within options="header" table
                for cell in rows[0]:
                    lines.append(f"| {cell}")
                lines.append("")
                for row in rows[1:]:
                    for cell in row:
                        lines.append(f"| {cell}")
                    lines.append("")
            else:
                for row in rows:
                    for cell in row:
                        lines.append(f"| {cell}")
                    lines.append("")
        lines += ["|===", ""]

    elif b.kind == KIND_FIGURE:
        img = b.extra.get("image_file", "")
        cap = b.extra.get("caption", b.text)
        if img:
            lines += ["", f".{cap}", f"image::../images/{img}[]", ""]
        else:
            lines += ["", f"// figure: {cap}", ""]

    elif b.kind == KIND_KEY:
        lines += ["", f"*Key:* {b.text}", ""]

    return lines


def write_metanorma(
    blocks: list[DocBlock], meta: dict, output_dir: Path, docx_path: Path, overwrite: bool = False
) -> dict:
    """Write a Metanorma AsciiDoc project to output_dir."""
    if output_dir.exists() and not overwrite:
        raise FileExistsError(f"Output directory already exists: {output_dir}")
    output_dir.mkdir(parents=True, exist_ok=True)

    sections_dir = output_dir / "sections"
    sections_dir.mkdir(exist_ok=True)

    images_dir = output_dir / "images"
    _extract_images(docx_path, images_dir)

    sections = _split_sections(blocks)
    include_lines: list[str] = []
    written: list[str] = []

    title = meta.get("title", "Untitled Specification")
    docnum = meta.get("docnumber", "")
    partnum = meta.get("partnumber", "")
    date = meta.get("date", "")

    for sec in sections:
        slug_part = re.sub(r"[^a-z0-9]+", "-", sec.title.lower())[:40].strip("-")
        fname = f"{sec.prefix}-{slug_part}.adoc"
        include_lines.append(f"include::sections/{fname}[]")
        written.append(fname)

        out_lines: list[str] = []

        if (ann_b := next((b for b in sec.blocks if b.kind == KIND_ANNEX_H), None)) is not None:
            marker = ann_b.extra.get("marker", "normative")
            slug = _slugify(ann_b.text)
            out_lines += [
                f"[appendix,obligation={marker}]",
                f"[[{slug}]]",
                f"== {ann_b.text}",
                "",
            ]
            for b in sec.blocks:
                if b.kind != KIND_ANNEX_H:
                    out_lines += _render_adoc_block(b, in_annex=True)
        else:
            if sec.prefix != "00":
                slug = _slugify(sec.title)
                out_lines += [f"[[{slug}]]", f"== {sec.title}", ""]
            for b in sec.blocks:
                out_lines += _render_adoc_block(b)

        (sections_dir / fname).write_text("\n".join(out_lines), encoding="utf-8")

    # Split title into intro/main parts for Metanorma structured title attributes.
    # ISO titles often follow: "Information technology — Main title — Part N: Part title"
    # Split on the LAST em-dash (` — `) to get intro + main.
    _EM = " \u2014 "
    if _EM in title:
        _split_pos = title.rfind(_EM)
        _title_intro = title[:_split_pos].strip()
        _title_main = title[_split_pos + len(_EM) :].strip()
    else:
        _title_intro = ""
        _title_main = title

    # Main document.adoc
    doc_lines = [
        f"= {title}",
        f":docnumber: {docnum}",
        f":partnumber: {partnum}",
        ":edition: 3",
    ]
    if date:
        doc_lines.append(f":revdate: {date}")
    if _title_intro:
        doc_lines.append(f":title-intro-en: {_title_intro}")
    doc_lines.append(f":title-main-en: {_title_main}")
    doc_lines += (
        [
            ":language: en",
            ":doctype: international-standard",
            ":copyright-year: " + (date[:4] if date else ""),
            ":mn-document-class: iso",
            ":mn-output-extensions: xml,html,pdf",
            "",
        ]
        + include_lines
        + [""]
    )

    (output_dir / "document.adoc").write_text("\n".join(doc_lines), encoding="utf-8")

    return {"sections": written, "warnings": [], "output_dir": output_dir}


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def convert_docx(
    docx_path: Path,
    output_dir: Path,
    fmt: str = "bikeshed",
    overwrite: bool = False,
    style_overrides: dict[str, str] | None = None,
) -> dict:
    """Convert an ISO Word file to Bikeshed or Metanorma format.

    Args:
        docx_path: Path to the .docx file.
        output_dir: Root output directory.
        fmt: ``"bikeshed"`` or ``"metanorma"``.
        overwrite: Allow overwriting existing output.
        style_overrides: Optional dict mapping Word style names to KIND_* constants.
            Entries with value ``"skip"`` suppress those paragraphs from output.
    """
    # Parse without extracting images here; each writer extracts once into the
    # final destination directory, then patches block image_file fields.
    blocks, meta = parse_docx(docx_path, images_dir=None, style_overrides=style_overrides)

    if fmt == "metanorma":
        result = write_metanorma(blocks, meta, output_dir, docx_path, overwrite=overwrite)
    else:
        result = write_bikeshed(blocks, meta, output_dir, docx_path, overwrite=overwrite)

    return result
