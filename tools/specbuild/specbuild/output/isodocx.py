"""ISO-styled Word document export.

Extends the base DOCX export pipeline with ISO-specific cover page,
heading styles, annex labeling, and bibliography formatting.
"""

from __future__ import annotations

import html as _html
import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from bs4 import BeautifulSoup

    from specbuild.standards.flavors import FlavorSpec


# ISO heading style mapping — keys are HTML tag names or section class names,
# values are the Word paragraph style names expected in the reference .docx template.
_ISO_HEADING_STYLES: dict[str, str] = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 5",
    "h6": "Heading 6",
    # Special ISO sections (matched by parent section class or heading id)
    "foreword": "Heading Foreword",
    "introduction": "Heading Introduction",
    "scope": "Heading Scope",
    "terms": "Heading Terms",
    "annex": "Heading Annex",
}


def generate_iso_docx(
    html_path: Path,
    output_path: Path,
    *,
    metadata: dict[str, str],
    flavor: FlavorSpec | None = None,
    reference_doc: Path | None = None,
    soup: BeautifulSoup | None = None,
) -> Path | None:
    """Generate an ISO-styled Word document.

    Extends the base DOCX export with:
    - ISO cover page with document number, stage, committee
    - ISO-specific heading styles (Clause prefix)
    - Proper annex handling with normative/informative labels
    - Copyright notice

    Args:
        html_path: Path to compiled index.html.
        output_path: Destination .docx path.
        metadata: Resolved standards metadata dict.
        flavor: Active standards flavor.
        reference_doc: Optional Word reference template.
        soup: Pre-parsed BeautifulSoup tree (avoids re-reading disk if provided).

    Returns:
        Path to generated DOCX file, or ``None`` on failure.
    """
    if soup is None:
        try:
            from specbuild.utils import read_html

            soup = read_html(html_path)
        except Exception:
            logging.error("Failed to read HTML for ISO DOCX export")
            return None

    _preprocess_for_iso(soup, metadata, flavor)

    preprocessed_html = html_path.parent / "_iso_preprocessed.html"
    from specbuild.utils import write_html

    write_html(preprocessed_html, soup)

    try:
        from specbuild.output.docxexport import generate_docx

        result = generate_docx(
            preprocessed_html,
            output_path,
            reference_doc=reference_doc,
            title=metadata.get("title_main", ""),
            branch="",
            sha="",
            date="",
            page_size="a4",
        )
    except Exception:
        logging.error("ISO DOCX generation failed", exc_info=True)
        result = None
    finally:
        preprocessed_html.unlink(missing_ok=True)

    if result:
        _postprocess_iso_docx(result, metadata, flavor)
        logging.info(f"ISO DOCX written to {output_path}")

    return result


def _preprocess_for_iso(
    soup: BeautifulSoup,
    metadata: dict[str, str],
    flavor: FlavorSpec | None,
) -> None:
    """Modify soup before DOCX conversion for ISO formatting."""
    from specbuild.enhancements.isonumbering import renumber_annexes_soup, renumber_clauses_soup

    _inject_iso_cover_info(soup, metadata, flavor)
    if flavor is not None:
        renumber_clauses_soup(soup, flavor)
        renumber_annexes_soup(soup, flavor)


def _inject_iso_cover_info(
    soup: BeautifulSoup,
    metadata: dict[str, str],
    flavor: FlavorSpec | None,
) -> None:
    """Inject ISO cover page information at the start of the document."""
    from bs4 import BeautifulSoup as BS
    from bs4 import NavigableString

    body = soup.find("body")
    if body is None:
        return

    docnumber = metadata.get("docnumber", "")
    partnumber = metadata.get("partnumber", "")
    stage = metadata.get("stage", "")
    title_main = metadata.get("title_main", "")
    title_part = metadata.get("title_part", "")
    committee = metadata.get("technical_committee", "")
    copyright_year = metadata.get("copyright_year", "")
    display_name = flavor.display_name if flavor else "ISO/IEC"

    doc_id = f"{display_name} {docnumber}"
    if partnumber:
        doc_id += f"-{partnumber}"

    from specbuild.standards.metadata import stage_display_name

    stage_text = stage_display_name(stage) if stage else ""

    _e = _html.escape
    cover_html = f"""<div class="iso-cover-page">
  <p class="iso-doc-id"><strong>{_e(doc_id)}</strong></p>
"""
    if stage_text:
        cover_html += f'  <p class="iso-stage">{_e(stage_text)}</p>\n'
    if title_main:
        cover_html += f'  <p class="iso-title"><strong>{_e(title_main)}</strong></p>\n'
    if title_part:
        cover_html += f'  <p class="iso-title-part">{_e(title_part)}</p>\n'
    if committee:
        cover_html += f'  <p class="iso-committee">{_e(committee)}</p>\n'
    if copyright_year:
        cover_html += (
            f'  <p class="iso-copyright">&copy; {_e(display_name)} {_e(copyright_year)}</p>\n'
        )
    cover_html += "</div>\n"

    _cover_doc = BS(cover_html, "html.parser")
    fragment = _cover_doc.find("div", class_="iso-cover-page") or next(
        (c for c in (_cover_doc.body or _cover_doc).children if getattr(c, "name", None)),
        None,
    )
    if fragment is None:
        logging.warning("ISO cover HTML produced no insertable element")
        return
    first_child = body.find()
    if first_child:
        first_child.insert_before(fragment)
        first_child.insert_before(NavigableString("\n"))
    else:
        body.append(fragment)


def _postprocess_iso_docx(
    docx_path: Path,
    metadata: dict[str, str],
    flavor: FlavorSpec | None,
) -> None:
    """Post-process the generated DOCX with ISO-specific metadata."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except ImportError:
        logging.debug("python-docx not available; skipping ISO DOCX post-processing")
        return

    try:
        doc = Document(str(docx_path))
    except Exception:
        logging.warning("Failed to open DOCX for ISO post-processing")
        return

    # --- Document metadata ---
    core = doc.core_properties
    title_main = metadata.get("title_main", "")
    if title_main and not core.title:
        core.title = title_main

    docnumber = metadata.get("docnumber", "")
    if docnumber:
        core.subject = docnumber

    committee = metadata.get("technical_committee", "")
    if committee and not core.author:
        core.author = committee

    # --- w:updateFields so TOC/cross-refs auto-update on open ---
    _inject_update_fields(doc)

    # --- Build the document identifier string for the running header ---
    display_name = flavor.display_name if flavor else "ISO/IEC"
    doc_id_parts = [display_name]
    if docnumber:
        doc_id_parts.append(f" {docnumber}")
        partnumber = metadata.get("partnumber", "")
        if partnumber:
            doc_id_parts.append(f"-{partnumber}")
    edition_year = metadata.get("edition_year", metadata.get("copyright_year", ""))
    if edition_year:
        doc_id_parts.append(f":{edition_year}")
    amendment = metadata.get("amendment_number", "")
    if amendment:
        doc_id_parts.append(f" Amendment {amendment}")
    header_text = "".join(doc_id_parts)

    # --- A4 page setup with ISO Directives Part 2 margins + running header ---
    for section in doc.sections:
        # A4: 210 mm × 297 mm
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        # ISO Directives Part 2: top/bottom 25 mm, left/right 20 mm
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

        # Right-aligned running header with document identifier
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # Clear any existing runs by removing their underlying XML elements
        for run in list(header_para.runs):
            run._r.getparent().remove(run._r)
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run(header_text)
        run.font.size = Pt(9)
        run.font.italic = True

    _apply_iso_heading_styles(doc)

    try:
        doc.save(str(docx_path))
    except Exception:
        logging.warning("Failed to save ISO DOCX post-processing changes")


def _inject_update_fields(doc: object) -> None:
    """Add ``<w:updateFields w:val="1"/>`` to document settings.

    This causes Word to prompt for a field update (TOC, page numbers,
    cross-references) the first time the document is opened.
    """
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return

    try:
        from specbuild.output.docxexport import _make_oxml_element
    except Exception:
        return

    try:
        settings = doc.settings.element  # type: ignore[union-attr]
        existing = settings.find(qn("w:updateFields"))
        if existing is not None:
            settings.remove(existing)
        uf = _make_oxml_element("w:updateFields", {"w:val": "true"})
        settings.append(uf)
    except Exception:
        logging.debug("Could not inject w:updateFields into ISO DOCX", exc_info=True)


def _apply_iso_heading_styles(doc: object) -> None:
    """Apply ISO heading styles to paragraphs in a python-docx Document.

    Remaps generic ``Heading N`` styles and special section headings to the
    ISO-specific style names defined in ``_ISO_HEADING_STYLES``.  Only styles
    that already exist in the document template are applied; unknown styles are
    left unchanged to avoid Word errors.

    Args:
        doc: A ``docx.Document`` instance (python-docx).
    """
    try:
        available_styles: set[str] = {s.name for s in doc.styles}
    except Exception:
        logging.debug("Could not enumerate DOCX styles; skipping ISO heading style application")
        return

    # Pandoc generates "Heading N" for <hN>; remap only where the ISO template style differs.
    _pandoc_style = {f"h{i}": f"Heading {i}" for i in range(1, 7)}
    style_remap: dict[str, str] = {}
    for html_tag, iso_style in _ISO_HEADING_STYLES.items():
        pandoc_style = _pandoc_style.get(html_tag)
        if pandoc_style and iso_style != pandoc_style and iso_style in available_styles:
            style_remap[pandoc_style] = iso_style

    if not style_remap:
        logging.debug(
            "No ISO heading style remapping needed (styles already match or template lacks ISO styles)"
        )
        return

    changed = 0
    try:
        for para in doc.paragraphs:
            if para.style and para.style.name in style_remap:
                para.style = doc.styles[style_remap[para.style.name]]
                changed += 1
    except Exception:
        logging.debug("Error applying ISO heading styles", exc_info=True)
        return

    if changed:
        logging.info(f"Applied ISO heading styles to {changed} paragraph(s)")
