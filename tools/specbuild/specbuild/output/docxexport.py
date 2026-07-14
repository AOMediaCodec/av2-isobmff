"""Word document (DOCX) export for compiled specifications.

Converts the compiled HTML specification into a native Microsoft Word document
using python-docx only — no external binaries. The compiled HTML DOM (already
enriched by SpecBuild with clause/figure/table numbering and resolved
cross-references) is walked directly into Word paragraphs, tables, figures,
lists and inline runs, then post-processed (metadata, headers/footers, page
setup) — the same enrich-once / serialize-many model used by the STS and IsoDoc
exporters.

Requirements:

- **python-docx** (``pip install python-docx``) — the only dependency.

Usage::

    python compile.py --docx                   # basic DOCX export
    python compile.py --docx --docx-template reference.docx  # with template
"""

from __future__ import annotations

import logging
import re
import shutil
import tempfile
from pathlib import Path

from specbuild.config import CONFIG
from specbuild.utils import HEADING_TAGS

# ---------------------------------------------------------------------------
# DOCX styling constants
# ---------------------------------------------------------------------------

_BORDER_COLOR = "999999"
_HEADER_SHADING = "E8E8E8"  # Light grey
_ALT_ROW_SHADING = "F5F5F5"  # Very light grey
_DATA_BORDER_COLOR = "BBBBBB"  # Lighter grey for data tables
_MONO_FONT = "Courier New"


# ---------------------------------------------------------------------------
# DOM / OXML helpers
# ---------------------------------------------------------------------------


def _wrap_children(parent, wrapper_tag) -> None:
    """Move all children of *parent* into *wrapper_tag* and append it.

    Handles both NavigableString and Tag nodes safely, avoiding
    double-escaping of HTML entities.
    """
    children = list(parent.children)
    parent.clear()
    for child in children:
        wrapper_tag.append(child.extract() if hasattr(child, "extract") else child)
    parent.append(wrapper_tag)


def _set_table_full_width(table, qn) -> None:
    """Set a Word table to 100% page width."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = _make_oxml_element("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")  # 100% in 50ths-of-a-percent


def _set_table_borders(
    table,
    qn,
    *,
    color: str = _BORDER_COLOR,
    edges: tuple[str, ...] = ("top", "left", "bottom", "right", "insideH", "insideV"),
) -> None:
    """Set single-line borders on a Word table."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = _make_oxml_element("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in edges:
        border = _make_oxml_element(
            f"w:{edge}",
            {"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": color},
        )
        existing = tbl_borders.find(qn(f"w:{edge}"))
        if existing is not None:
            tbl_borders.remove(existing)
        tbl_borders.append(border)


def _set_cell_shading(cell, qn, fill: str) -> None:
    """Apply background shading to a Word table cell."""
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = _make_oxml_element("w:tcPr")
        cell._tc.insert(0, tc_pr)
    existing_shd = tc_pr.find(qn("w:shd"))
    if existing_shd is not None:
        tc_pr.remove(existing_shd)
    tc_pr.append(
        _make_oxml_element(
            "w:shd",
            {"w:val": "clear", "w:fill": fill, "w:color": "auto"},
        )
    )


# ---------------------------------------------------------------------------
# Pandoc availability check
# ---------------------------------------------------------------------------


def _check_python_docx() -> bool:
    """Return True if python-docx is importable."""
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        logging.warning(
            "python-docx not available; DOCX post-processing will be skipped. "
            "Install with: pip install python-docx"
        )
        return False


# ---------------------------------------------------------------------------
# HTML preprocessing for DOCX
# ---------------------------------------------------------------------------

# Elements and attributes that are web-only and should be removed or
# simplified before DOCX conversion.
_WEB_ONLY_IDS = {
    "table-of-changes",
    "links-banner",
    "pwa-install-banner",
    "back-to-toc",
    "dark-mode-toggle",
}

# CSS style blocks injected by the build system (not useful in DOCX)
_STRIP_STYLE_IDS = {
    "line-anchors-css",
    "syntax-tooltips-css",
    "figure-table-tooltips-css",
    "dark-mode-css",
    "change-bars-css",
    "toc-bold-primary-only-css",
    "table-of-changes-css",
    "watermark-css",
    "cover-page-css",
    "page-numbering-css",
    "section-headers-css",
    "keyword-highlighting-css",
}


def preprocess_html_for_docx(html_path: Path, output_path: Path) -> None:
    """Preprocess compiled HTML for better DOCX conversion.

    Strips web-only elements (scripts, injected styles, navigation),
    fixes image paths to be relative to the HTML file, and simplifies
    markup that pandoc handles poorly.

    Args:
        html_path: Path to the compiled ``index.html``.
        output_path: Path to write the preprocessed HTML.
    """
    from specbuild.utils import get_bs4

    try:
        BS = get_bs4()
    except ImportError:
        # Fallback: copy as-is
        shutil.copy2(html_path, output_path)
        return

    html_text = html_path.read_text(encoding="utf-8")
    # Use html5lib to produce well-formed HTML — the html.parser output
    # can confuse pandoc when implicit closing tags are serialised.
    try:
        soup = BS(html_text, "html5lib")
    except Exception:
        logging.warning("html5lib parser not available, falling back to html.parser")
        soup = BS(html_text, "html.parser")

    changes = 0

    # Remove all <script> tags
    for script in soup.find_all("script"):
        script.decompose()
        changes += 1

    # Remove injected <style> blocks by known IDs
    for style in soup.find_all("style", id=True):
        if style["id"] in _STRIP_STYLE_IDS:
            style.decompose()
            changes += 1

    # Remove web-only elements by ID
    for elem_id in _WEB_ONLY_IDS:
        elem = soup.find(id=elem_id)
        if elem:
            elem.decompose()
            changes += 1

    # Remove navigation elements (back-to-toc links, etc.)
    for link in soup.find_all("a", class_="back-to-toc"):
        link.decompose()
        changes += 1

    # Remove data-tooltip attributes (noise in DOCX)
    for elem in soup.find_all(attrs={"data-tooltip": True}):
        del elem["data-tooltip"]
        changes += 1

    # Remove Bikeshed self-link anchors (pandoc renders href as text)
    for a in soup.find_all("a", class_="self-link"):
        a.decompose()
        changes += 1

    # Remove line-anchor wrapper spans (keep content)
    for span in soup.find_all("span", class_="code-line"):
        span.unwrap()
        changes += 1

    # Simplify change bars — remove the CSS class but keep content
    for elem in soup.find_all(class_="changed"):
        elem["class"] = [c for c in elem.get("class", []) if c != "changed"]
        if not elem.get("class"):
            del elem["class"]
        changes += 1

    # Fix watermark overlay (remove entirely for DOCX)
    watermark = soup.find(class_="watermark-overlay")
    if watermark:
        watermark.decompose()
        changes += 1

    # ── Strip non-essential IDs to reduce Word bookmarks ──
    # Pandoc converts every HTML id= into a Word bookmark.  Keep only
    # IDs on headings, captions, figcaptions, and definitions (useful for
    # document navigation / cross-references).  Strip everything else.
    _KEEP_ID_TAGS = HEADING_TAGS | {"caption", "figcaption", "dfn"}
    for elem in soup.find_all(id=True):
        if elem.name not in _KEEP_ID_TAGS:
            del elem["id"]
            changes += 1

    # ── Heading bookmark anchors ──
    # Pandoc uses heading IDs for TOC generation but does NOT create
    # explicit Word bookmarks on heading paragraphs.  Internal hyperlinks
    # (§ cross-references) need bookmarks to resolve.  Fix: insert a
    # <span id="..."> inside each heading so pandoc creates a bookmark.
    for heading in soup.find_all(HEADING_TAGS):
        heading_id = heading.get("id")
        if heading_id:
            anchor_span = soup.new_tag("span", id=heading_id)
            anchor_span.string = ""  # empty anchor
            heading.insert(0, anchor_span)
            del heading["id"]  # remove from heading to avoid duplicate
            changes += 1

    # ── Definition terms: wrap in <strong> for bold ──
    # Bikeshed <dfn> elements are bold via CSS (font-weight: bolder)
    # but pandoc ignores CSS, rendering them as plain text.
    for dfn in soup.find_all("dfn"):
        # Skip if already inside <strong> or <b>
        if dfn.find_parent("strong") or dfn.find_parent("b"):
            continue
        strong = soup.new_tag("strong")
        dfn.wrap(strong)
        changes += 1

    # ── Definition list terms (<dt>): wrap in <strong> for bold ──
    # CSS styles <dt> as bold (font-weight: bold) for bibliography
    # entries and other definition lists.
    for dt in soup.find_all("dt"):
        if dt.find("strong") or dt.find("b"):
            continue  # already has bold content
        strong = soup.new_tag("strong")
        _wrap_children(dt, strong)
        changes += 1

    # ── Note markers: wrap "Note:" prefix in <strong> ──
    # Notes have <span class="marker">Note:</span> styled via CSS.
    for marker in soup.find_all("span", class_="marker"):
        strong = soup.new_tag("strong")
        marker.wrap(strong)
        changes += 1

    # ── SDL syntax tables: convert CSS-dependent styling to inline HTML ──
    # Pandoc ignores CSS classes, so we need to:
    #   1. Convert padding-left indentation to non-breaking spaces
    #   2. Wrap variable names in <strong> (bold via CSS class in HTML)
    #   3. Wrap code content in <code> for monospace rendering
    for sdl_table in soup.find_all("table", class_="sdl-syntax-table"):
        # Mark the table for post-processing identification
        sdl_table["data-sdl"] = "true"

        for td in sdl_table.find_all("td"):
            classes = td.get("class", [])
            # Convert padding-left on inner spans to non-breaking spaces
            for span in td.find_all("span", style=True):
                style = span.get("style", "")
                match = re.search(r"padding-left:\s*([\d.]+)em", style)
                if match:
                    em_val = float(match.group(1))
                    # ~2 non-breaking spaces per em for readable indentation
                    nbsp_count = round(em_val * 2)
                    nbsp_prefix = "\u00a0" * nbsp_count
                    # Prepend spaces and unwrap the span
                    if span.string:
                        span.string = nbsp_prefix + span.string
                    else:
                        span.insert(0, nbsp_prefix)
                    del span["style"]
                    span.unwrap()
                    changes += 1

            # Wrap variable names (sdl-var-with-descriptor) in <strong>
            if "sdl-var-with-descriptor" in classes:
                code = soup.new_tag("code")
                _wrap_children(td, code)
                strong = soup.new_tag("strong")
                td.clear()
                strong.append(code)
                td.append(strong)
                changes += 1
            # Wrap plain code lines in <code> for monospace
            elif "sdl-code" in classes:
                _wrap_children(td, soup.new_tag("code"))
                changes += 1
            # Wrap descriptor text in <code>
            elif "sdl-descriptor" in classes:
                if td.get_text(strip=True):
                    _wrap_children(td, soup.new_tag("code"))
                    changes += 1

        # Wrap header cells in <code> too
        for th in sdl_table.find_all("th"):
            classes = th.get("class", [])
            if "sdl-syntax-name" in classes:
                code = soup.new_tag("code")
                _wrap_children(th, code)
                strong = soup.new_tag("strong")
                th.clear()
                strong.append(code)
                th.append(strong)
                changes += 1

        # Remove data-original-syntax (bloats DOCX)
        if sdl_table.has_attr("data-original-syntax"):
            del sdl_table["data-original-syntax"]

    # ── Prevent adjacent SDL tables from merging ──
    # Pandoc merges consecutive tables with no content between them.
    # Insert a thin horizontal rule between adjacent SDL tables.
    sdl_tables = soup.find_all("table", class_="sdl-syntax-table")
    for i in range(len(sdl_tables) - 1):
        tbl = sdl_tables[i]
        nxt = sdl_tables[i + 1]
        # Check if there's any block element between them
        has_separator = False
        sibling = tbl.next_sibling
        while sibling and sibling is not nxt:
            if hasattr(sibling, "name") and sibling.name in (
                "p",
                "div",
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "hr",
            ):
                has_separator = True
                break
            sibling = sibling.next_sibling
        if not has_separator:
            # Insert a thin spacer paragraph
            spacer = soup.new_tag("p")
            spacer.string = "\u00a0"  # non-breaking space
            tbl.insert_after(spacer)
            changes += 1

    # ── Code-block tables: wrap content in <code> for monospace ──
    for code_table in soup.find_all("table", class_="code-table"):
        for td in code_table.find_all("td"):
            if td.get_text(strip=True):
                _wrap_children(td, soup.new_tag("code"))
                changes += 1

    # Simplify equation wrappers so pandoc sees raw LaTeX math.
    # The build system wraps equations as:
    #   <div class="equation-wrapper"><p>$$...$$</p>
    #       <span class="equation-number">(5.1)</span></div>
    # Replace with a borderless two-column table: equation centered,
    # number right-aligned — the standard Word pattern for numbered
    # display equations.
    for wrapper in soup.find_all(class_="equation-wrapper"):
        eq_num = wrapper.find(class_="equation-number")
        num_text = eq_num.get_text(strip=True) if eq_num else ""
        if eq_num:
            eq_num.decompose()
        math_p = wrapper.find("p")
        math_content = math_p.decode_contents() if math_p else ""
        # Build a borderless table: | equation (center) | number (right) |
        tbl_html = (
            '<table style="width:100%;border:none;border-collapse:collapse">'
            "<tr>"
            f'<td style="text-align:center;border:none">{math_content}</td>'
            f'<td style="text-align:right;width:60px;border:none;vertical-align:middle">{num_text}</td>'
            "</tr></table>"
        )
        new_tag = soup.new_tag("div")
        new_tag.append(BS(tbl_html, "html.parser").find("table"))
        wrapper.replace_with(new_tag)
        new_tag.unwrap()
        changes += 1

    # Fix cover page — convert to simple heading content
    cover = soup.find(id="cover-page")
    if cover:
        cover.decompose()
        changes += 1

    # ── Footnote normalization ──
    # Bikeshed: <sup><a href="#fn-1" class="footnote-ref">1</a></sup>
    # Pandoc expects: <a href="#fn-1" class="footnote-ref" role="doc-noteref"><sup>1</sup></a>
    # Normalize structure so pandoc reliably extracts footnotes when reading HTML.
    for sup in soup.find_all("sup"):
        a = sup.find("a", class_="footnote-ref")
        if a is None:
            continue
        # Rewrap: move <sup> content inside <a>, then replace <sup> with <a>
        a_new = soup.new_tag("a")
        a_new["href"] = a.get("href", "")
        a_new["class"] = "footnote-ref"
        a_new["role"] = "doc-noteref"
        if a.get("id"):
            a_new["id"] = a["id"]
        inner_sup = soup.new_tag("sup")
        for child in list(a.children):
            inner_sup.append(child.extract() if hasattr(child, "extract") else child)
        a_new.append(inner_sup)
        sup.replace_with(a_new)
        changes += 1

    # Mark the footnotes section with ARIA role so pandoc treats it as endnotes
    for fn_section in soup.find_all(class_="footnotes"):
        if not fn_section.get("role"):
            fn_section["role"] = "doc-endnotes"
            changes += 1

    # Normalize return-link class: footnote-backref → footnote-back
    for back_a in soup.find_all("a", class_="footnote-backref"):
        classes = back_a.get("class", [])
        back_a["class"] = [c if c != "footnote-backref" else "footnote-back" for c in classes]
        if not back_a.get("role"):
            back_a["role"] = "doc-backlink"
        changes += 1

    logging.debug(f"DOCX preprocessing: {changes} modifications")
    output_path.write_text(str(soup), encoding="utf-8")


# ---------------------------------------------------------------------------
# Pandoc conversion
# ---------------------------------------------------------------------------


def _build_docx_native(
    html_path: Path,
    docx_path: Path,
    *,
    reference_doc: Path | None = None,
    title: str | None = None,
    toc: bool = True,
    toc_depth: int = 3,
    resource_path: Path | None = None,
) -> bool:
    """Build a DOCX natively from the (preprocessed) HTML DOM with python-docx.

    Walks the enriched HTML — which already carries SpecBuild's clause/figure/
    table numbering and resolved cross-references — into Word headings,
    paragraphs, tables, figures, lists and inline runs. No external tools.

    Args:
        html_path: Input (preprocessed) HTML file.
        docx_path: Output DOCX file.
        reference_doc: Optional Word template supplying base styles.
        title: Document title (emitted as the Title paragraph).
        toc: Whether to insert a Table of Contents field.
        toc_depth: TOC heading depth.
        resource_path: Directory for resolving relative image paths.

    Returns:
        True on success, False if python-docx is unavailable or the build fails.
    """
    try:
        from docx import Document
    except ImportError:
        logging.error("python-docx not available; cannot generate DOCX. Install python-docx.")
        return False

    from specbuild.utils import get_bs4

    try:
        BS = get_bs4()
        soup = BS(html_path.read_text(encoding="utf-8"), "html.parser")
    except Exception:
        logging.error("Failed to parse HTML for DOCX build", exc_info=True)
        return False

    try:
        doc = Document(str(reference_doc)) if reference_doc and reference_doc.exists() else Document()
        writer = _DocxWriter(doc, base_dir=resource_path or html_path.parent)
        writer.build(soup, title=title, toc=toc, toc_depth=toc_depth)
        doc.save(str(docx_path))
        return True
    except Exception:
        logging.error("Native DOCX build failed", exc_info=True)
        return False


# Block-level HTML tags handled directly by the writer (others recurse).
_BLOCK_TAGS = frozenset(
    {
        "section",
        "div",
        "p",
        "ul",
        "ol",
        "dl",
        "table",
        "figure",
        "pre",
        "blockquote",
        "aside",
        "details",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
    }
)
_HEADING_TAGS = frozenset({"h1", "h2", "h3", "h4", "h5", "h6"})
# Inline tags producing character formatting.
_BOLD_TAGS = frozenset({"strong", "b", "dfn"})
_ITALIC_TAGS = frozenset({"em", "i", "var", "cite"})
_MONO_TAGS = frozenset({"code", "tt", "kbd", "samp"})


class _DocxWriter:
    """Render an enriched-HTML soup into a python-docx ``Document``."""

    def __init__(self, doc, base_dir: Path) -> None:
        self.doc = doc
        self.base_dir = base_dir

    # -- entry ---------------------------------------------------------------
    def build(self, soup, *, title: str | None, toc: bool, toc_depth: int) -> None:
        if title:
            try:
                self.doc.add_heading(title, level=0)
            except Exception:
                self.doc.add_paragraph(title)
        if toc:
            self._add_toc(toc_depth)
        body = soup.find("main") or soup.find("body") or soup
        self._render_blocks(body)

    # -- block dispatch ------------------------------------------------------
    def _render_blocks(self, container) -> None:
        from bs4 import NavigableString, Tag

        for child in getattr(container, "children", []):
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    self.doc.add_paragraph(text)
                continue
            if not isinstance(child, Tag):
                continue
            name = child.name
            if name in _HEADING_TAGS:
                self._render_heading(child)
            elif name == "p":
                self._render_paragraph(child)
            elif name in ("ul", "ol"):
                self._render_list(child, ordered=(name == "ol"), level=0)
            elif name == "dl":
                self._render_dl(child)
            elif name == "table":
                self._render_table(child)
            elif name == "figure":
                self._render_figure(child)
            elif name == "pre":
                self._render_pre(child)
            elif name in ("section", "div", "blockquote", "aside", "details"):
                # Transparent wrappers — recurse into their children.
                self._render_blocks(child)
            # Unknown/inline-only blocks: ignore (scripts/styles stripped earlier).

    # -- headings ------------------------------------------------------------
    def _render_heading(self, el) -> None:
        level = int(el.name[1])
        text = el.get_text(" ", strip=True)
        if not text:
            return
        # Word built-in heading styles go to level 9; clamp.
        para = self.doc.add_heading(level=min(level, 9))
        self._render_inline(para, el)
        # add_heading() created an empty run-less paragraph; if inline produced
        # nothing, fall back to plain text.
        if not para.runs:
            para.add_run(text)

    # -- paragraphs ----------------------------------------------------------
    def _render_paragraph(self, el) -> None:
        if not el.get_text(strip=True) and not el.find("img"):
            return
        para = self.doc.add_paragraph()
        self._render_inline(para, el)

    # -- lists ---------------------------------------------------------------
    def _render_list(self, el, *, ordered: bool, level: int) -> None:
        from bs4 import Tag

        style = "List Number" if ordered else "List Bullet"
        if level > 0:
            style = f"{style} {level + 1}"
        for li in el.find_all("li", recursive=False):
            para = self.doc.add_paragraph(style=self._safe_style(style, ordered, level))
            self._render_inline(para, li, skip_nested_lists=True)
            for sub in li.find_all(["ul", "ol"], recursive=False):
                self._render_list(sub, ordered=(sub.name == "ol"), level=level + 1)
            # Block children inside <li> (e.g. nested <p>, <table>) — render after.
            for blk in li.find_all(["p", "table", "figure", "pre"], recursive=False):
                if blk.name == "p":
                    self._render_paragraph(blk)
                elif blk.name == "table":
                    self._render_table(blk)
                elif blk.name == "figure":
                    self._render_figure(blk)
                elif blk.name == "pre":
                    self._render_pre(blk)
            _ = Tag  # silence unused in some linters

    def _safe_style(self, style: str, ordered: bool, level: int):
        """Return *style* if present in the document, else a safe fallback."""
        try:
            _ = self.doc.styles[style]
            return style
        except KeyError:
            return "List Number" if ordered else "List Bullet"

    # -- definition lists ----------------------------------------------------
    def _render_dl(self, el) -> None:
        nodes = el.find_all(["dt", "dd"], recursive=False)
        pending_terms: list = []
        for node in nodes:
            if node.name == "dt":
                pending_terms.append(node)
            else:  # dd
                term_text = " ".join(t.get_text(" ", strip=True) for t in pending_terms).strip()
                pending_terms = []
                para = self.doc.add_paragraph()
                if term_text:
                    run = para.add_run(term_text + "  ")
                    run.bold = True
                self._render_inline(para, node)

    # -- preformatted --------------------------------------------------------
    def _render_pre(self, el) -> None:
        text = el.get_text()
        para = self.doc.add_paragraph(style=self._safe_style("No Spacing", False, 0))
        run = para.add_run(text)
        run.font.name = "Courier New"

    # -- figures -------------------------------------------------------------
    def _render_figure(self, el) -> None:
        from docx.shared import Inches

        img = el.find("img")
        if img is not None:
            src = img.get("src", "")
            if src and not src.startswith(("http://", "https://", "data:")):
                img_path = (self.base_dir / src).resolve()
                if img_path.exists():
                    try:
                        self.doc.add_picture(str(img_path), width=Inches(6.0))
                    except Exception:
                        logging.debug("Could not embed image %s", img_path)
        cap = el.find("figcaption")
        if cap is not None and cap.get_text(strip=True):
            para = self.doc.add_paragraph()
            run = para.add_run(cap.get_text(" ", strip=True))
            run.italic = True
            run.font.size = self._pt(9)

    # -- tables --------------------------------------------------------------
    def _render_table(self, el) -> None:
        cap = el.find("caption")
        if cap is not None and cap.get_text(strip=True):
            para = self.doc.add_paragraph()
            run = para.add_run(cap.get_text(" ", strip=True))
            run.bold = True
            run.font.size = self._pt(9)

        # Gather rows from thead+tbody (or the table directly), non-recursively.
        rows = []
        for grp in el.find_all(["thead", "tbody", "tfoot"], recursive=False):
            rows.extend(grp.find_all("tr", recursive=False))
        rows.extend(el.find_all("tr", recursive=False))
        if not rows:
            return
        ncols = max(len(r.find_all(["td", "th"], recursive=False)) for r in rows)
        if ncols == 0:
            return

        table = self.doc.add_table(rows=0, cols=ncols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for tr in rows:
            cells = tr.find_all(["td", "th"], recursive=False)
            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                if i >= ncols:
                    break
                tc = row_cells[i]
                # Reuse the auto-created empty paragraph.
                para = tc.paragraphs[0]
                self._render_inline(para, cell)
                if cell.name == "th":
                    for run in para.runs:
                        run.bold = True

    # -- inline runs ---------------------------------------------------------
    def _render_inline(self, para, el, *, skip_nested_lists: bool = False) -> None:
        self._emit_inline(para, el, bold=False, italic=False, mono=False,
                          skip_nested_lists=skip_nested_lists)

    def _emit_inline(self, para, node, *, bold, italic, mono, skip_nested_lists=False) -> None:
        from bs4 import NavigableString, Tag

        for child in getattr(node, "children", []):
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = para.add_run(text)
                    run.bold = bold or None
                    run.italic = italic or None
                    if mono:
                        run.font.name = "Courier New"
                continue
            if not isinstance(child, Tag):
                continue
            name = child.name
            if name in ("script", "style"):
                continue
            if skip_nested_lists and name in ("ul", "ol"):
                continue
            if name in ("p", "table", "figure", "pre") and skip_nested_lists:
                continue  # block children of <li> handled separately
            if name == "br":
                para.add_run().add_break()
                continue
            if name == "img":
                continue  # images handled at figure level
            if name == "a":
                self._emit_link(para, child, bold=bold, italic=italic, mono=mono)
                continue
            nb = bold or (name in _BOLD_TAGS)
            ni = italic or (name in _ITALIC_TAGS)
            nm = mono or (name in _MONO_TAGS)
            self._emit_inline(para, child, bold=nb, italic=ni, mono=nm,
                              skip_nested_lists=skip_nested_lists)

    def _emit_link(self, para, a, *, bold, italic, mono) -> None:
        """Emit an <a> as a hyperlink (internal bookmark or external URL)."""
        href = a.get("href", "")
        text = a.get_text(" ", strip=True)
        if not text:
            return
        try:
            if href.startswith("#"):
                self._add_internal_hyperlink(para, href[1:], text)
                return
            if href.startswith(("http://", "https://", "mailto:")):
                self._add_external_hyperlink(para, href, text)
                return
        except Exception:
            pass  # fall back to styled text
        run = para.add_run(text)
        run.bold = bold or None
        run.italic = italic or None

    def _add_internal_hyperlink(self, para, anchor: str, text: str) -> None:
        from docx.oxml.ns import qn

        hyperlink = _make_oxml_element("w:hyperlink", {"w:anchor": anchor})
        run = _make_oxml_element("w:r")
        rpr = _make_oxml_element("w:rPr")
        rstyle = _make_oxml_element("w:rStyle", {"w:val": "Hyperlink"})
        rpr.append(rstyle)
        run.append(rpr)
        wt = _make_oxml_element("w:t")
        wt.text = text
        wt.set(qn("xml:space"), "preserve")
        run.append(wt)
        hyperlink.append(run)
        para._p.append(hyperlink)

    def _add_external_hyperlink(self, para, url: str, text: str) -> None:
        part = para.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        from docx.oxml.ns import qn

        hyperlink = _make_oxml_element("w:hyperlink", {qn("r:id"): r_id})
        run = _make_oxml_element("w:r")
        rpr = _make_oxml_element("w:rPr")
        rstyle = _make_oxml_element("w:rStyle", {"w:val": "Hyperlink"})
        rpr.append(rstyle)
        run.append(rpr)
        wt = _make_oxml_element("w:t")
        wt.text = text
        wt.set(qn("xml:space"), "preserve")
        run.append(wt)
        hyperlink.append(run)
        para._p.append(hyperlink)

    # -- table of contents ---------------------------------------------------
    def _add_toc(self, depth: int) -> None:
        para = self.doc.add_paragraph()
        run = para.add_run()
        fld_begin = _make_oxml_element("w:fldChar", {"w:fldCharType": "begin"})
        instr = _make_oxml_element("w:instrText", {"xml:space": "preserve"})
        instr.text = f'TOC \\o "1-{depth}" \\h \\z \\u'
        fld_sep = _make_oxml_element("w:fldChar", {"w:fldCharType": "separate"})
        fld_end = _make_oxml_element("w:fldChar", {"w:fldCharType": "end"})
        r = run._r
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_sep)
        r.append(fld_end)

    def _pt(self, size: int):
        from docx.shared import Pt

        return Pt(size)



# ---------------------------------------------------------------------------
# Post-processing with python-docx
# ---------------------------------------------------------------------------


def _postprocess_docx(
    docx_path: Path,
    *,
    title: str | None = None,
    branch: str = "",
    sha: str = "",
    date: str = "",
    organization: str | None = None,
    page_size: str = "letter",
) -> None:
    """Post-process the DOCX file with python-docx.

    Adds document metadata, sets page layout, and inserts headers/footers.

    Args:
        docx_path: Path to the DOCX file to modify.
        title: Document title for metadata.
        branch: Git branch name.
        sha: Git commit SHA.
        date: Build date string.
        organization: Organization name for the footer.
        page_size: Page size: "letter", "a4", or "legal".
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt
    except ImportError:
        logging.warning("python-docx not available; skipping DOCX post-processing")
        return

    doc = Document(str(docx_path))

    # --- Document metadata ---
    core = doc.core_properties
    if title:
        core.title = title
    core.author = organization or CONFIG.spec_full_name
    if sha:
        core.revision = 1
        core.comments = f"Built from {branch}@{sha} on {date}"
    core.language = "en-US"

    # --- Auto-update fields on open ---
    # The TOC field and PAGE/NUMPAGES footer fields need to be updated
    # when the document opens.  Word will prompt "Do you want to update
    # the fields in this document?" — answering Yes populates the TOC
    # with page numbers and resolves all field codes.
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is not None:
        settings.remove(update_fields)
    uf = _make_oxml_element("w:updateFields", {"w:val": "true"})
    settings.append(uf)

    # --- Page setup ---
    _PAGE_SIZES = {
        "letter": (Inches(8.5), Inches(11)),
        "a4": (Cm(21.0), Cm(29.7)),
        "legal": (Inches(8.5), Inches(14)),
    }
    width, height = _PAGE_SIZES.get(page_size, _PAGE_SIZES["letter"])

    for section in doc.sections:
        section.page_width = width
        section.page_height = height
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # --- Header ---
        header = section.header
        header.is_linked_to_previous = False
        if not header.paragraphs or not header.paragraphs[0].text:
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run(title or CONFIG.spec_full_name)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = _docx_color(0x99, 0x99, 0x99)

        # --- Footer ---
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs or not footer.paragraphs[0].text:
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Version info
            if sha:
                version_run = footer_para.add_run(f"{branch}@{sha[:8]}  —  {date}    ")
                version_run.font.size = Pt(8)
                version_run.font.color.rgb = _docx_color(0x99, 0x99, 0x99)

            # Page number field
            _add_page_number_field(footer_para)

    # --- Table styling: single-pass classification and styling ---
    sdl_count = code_count = eq_count = data_count = 0
    for table in doc.tables:
        ncols = len(table.columns)
        nrows = len(table.rows)
        if not nrows:
            continue

        # Equation tables: 2-col, 1-row, right cell matches "(X.Y)"
        if ncols == 2 and nrows == 1:
            if len(table.rows[0].cells) > 1:
                right_text = table.rows[0].cells[1].text.strip()
                if re.match(r"^\([\w.]+\)$", right_text):
                    _set_table_full_width(table, qn)
                    eq_count += 1
                    continue

        # SDL tables: 2-col with "Descriptor" header
        if ncols == 2 and nrows > 1:
            if len(table.rows[0].cells) > 1:
                header_right = table.rows[0].cells[1].text.strip()
                if header_right == "Descriptor":
                    _style_word_table(table, qn, has_header=True)
                    sdl_count += 1
                    continue

        # Code-block tables: single-column
        if ncols == 1:
            _style_word_table(table, qn, has_header=False)
            code_count += 1
            continue

        # Data tables: everything else gets light borders + header shading
        _style_data_table(table, qn)
        data_count += 1

    if sdl_count:
        logging.info(f"Styled {sdl_count} SDL syntax tables")
    if code_count:
        logging.info(f"Styled {code_count} code-block tables")
    if data_count:
        logging.info(f"Styled {data_count} data tables")

    doc.save(str(docx_path))
    logging.info("DOCX post-processing complete: metadata, headers, footers, page setup")


def _style_word_table(table, qn, *, has_header: bool) -> None:
    """Apply specification-quality styling to an SDL or code-block table.

    Adds borders, monospace font, full page width, and alternating row
    shading.  When *has_header* is True, the first row gets header
    shading, bold text, larger font, and ``tblHeader`` is removed
    (prevents confusing repeated headers in Word).

    Args:
        table: A python-docx Table object.
        qn: The ``docx.oxml.ns.qn`` namespace resolver.
        has_header: Whether the first row is a header row.
    """
    from docx.shared import Pt

    font_size = Pt(8)
    header_font_size = Pt(9)
    edges = (
        ("top", "left", "bottom", "right", "insideH", "insideV")
        if has_header
        else ("top", "left", "bottom", "right", "insideH")
    )

    _set_table_full_width(table, qn)
    _set_table_borders(table, qn, edges=edges)

    # Remove "Repeat Header Rows" on all rows
    if has_header:
        for row in table.rows:
            tr_pr = row._tr.find(qn("w:trPr"))
            if tr_pr is not None:
                tbl_header = tr_pr.find(qn("w:tblHeader"))
                if tbl_header is not None:
                    tr_pr.remove(tbl_header)

    for row_idx, row in enumerate(table.rows):
        is_header_row = has_header and row_idx == 0
        is_alt = row_idx % 2 == 0 and not is_header_row

        for cell in row.cells:
            if is_header_row:
                _set_cell_shading(cell, qn, _HEADER_SHADING)
            elif is_alt:
                _set_cell_shading(cell, qn, _ALT_ROW_SHADING)

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = _MONO_FONT
                    run.font.size = header_font_size if is_header_row else font_size
                    if is_header_row:
                        run.font.bold = True


def _style_data_table(table, qn) -> None:
    """Add light borders and header shading to a data table.

    Skips tables that already have borders defined.
    """
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return
    # Only add borders if none exist
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is not None:
        has_border = any(
            tbl_borders.find(qn(f"w:{edge}")) is not None for edge in ("top", "bottom", "insideH")
        )
        if has_border:
            return
    _set_table_borders(table, qn, color=_DATA_BORDER_COLOR)

    # Header row shading
    if len(table.rows) > 1:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, qn, _HEADER_SHADING)


def _docx_color(r: int, g: int, b: int):
    """Create a python-docx RGBColor."""
    from docx.shared import RGBColor

    return RGBColor(r, g, b)


def _add_page_number_field(paragraph) -> None:
    """Insert a PAGE / NUMPAGES field code into a paragraph.

    Produces "Page X of Y" using Word field codes.
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt

    grey = _docx_color(0x99, 0x99, 0x99)

    def _add_field_run(field_name: str) -> None:
        """Append a Word field code run (e.g. PAGE, NUMPAGES) to the paragraph."""
        fld_begin = _make_oxml_element("w:fldChar", {"w:fldCharType": "begin"})
        instr = _make_oxml_element("w:instrText")
        instr.text = f" {field_name} "
        instr.set(qn("xml:space"), "preserve")
        fld_end = _make_oxml_element("w:fldChar", {"w:fldCharType": "end"})

        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = grey
        run._element.append(fld_begin)
        run._element.append(instr)
        run._element.append(fld_end)

    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = grey

    _add_field_run("PAGE")

    of_run = paragraph.add_run(" of ")
    of_run.font.size = Pt(8)
    of_run.font.color.rgb = grey

    _add_field_run("NUMPAGES")


def _make_oxml_element(tag: str, attrib: dict | None = None):
    """Create an OxmlElement with optional attributes."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    elem = OxmlElement(tag)
    if attrib:
        for key, val in attrib.items():
            elem.set(qn(key), val)
    return elem


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_docx(
    html_path: Path,
    output_path: Path,
    *,
    reference_doc: Path | None = None,
    title: str | None = None,
    branch: str = "",
    sha: str = "",
    date: str = "",
    organization: str | None = None,
    page_size: str = "letter",
    toc: bool = True,
    toc_depth: int = 3,
) -> Path | None:
    """Convert a compiled HTML specification to a Word document.

    The conversion pipeline:

    1. Preprocess the HTML (strip web-only elements, fix paths)
    2. Build the DOCX natively from the HTML DOM with python-docx
    3. Post-process with python-docx (metadata, headers, footers, page setup)

    Args:
        html_path: Path to the compiled ``index.html``.
        output_path: Desired output ``.docx`` path.
        reference_doc: Optional Word template file for styling.
        title: Document title (defaults to spec config title).
        branch: Git branch name for metadata.
        sha: Git commit SHA for metadata.
        date: Build date for metadata.
        organization: Organization name for footer.
        page_size: Page size: ``"letter"``, ``"a4"``, or ``"legal"``.
        toc: Whether to generate a table of contents.

    Returns:
        Path to the generated DOCX file, or ``None`` on failure.
    """
    if not _check_python_docx():
        logging.error("python-docx is required for DOCX export. Install with: pip install python-docx")
        return None

    if not html_path.exists():
        logging.error(f"HTML file not found: {html_path}")
        return None

    doc_title = title or CONFIG.spec_full_name

    logging.info(f"Generating Word document from {html_path.name} (native python-docx)")

    # Step 1: Preprocess HTML
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        preprocessed = tmp / "preprocessed.html"
        preprocess_html_for_docx(html_path, preprocessed)

        # Step 2: Build the DOCX natively from the (enriched) HTML DOM. Resolve
        # relative image paths against the original HTML location.
        tmp_docx = tmp / "output.docx"
        success = _build_docx_native(
            preprocessed,
            tmp_docx,
            reference_doc=reference_doc,
            title=doc_title,
            toc=toc,
            toc_depth=toc_depth,
            resource_path=html_path.parent,
        )
        if not success:
            return None

        # Step 3: Post-process with python-docx
        _postprocess_docx(
            tmp_docx,
            title=doc_title,
            branch=branch,
            sha=sha,
            date=date,
            organization=organization,
            page_size=page_size,
        )

        # Move to final location
        shutil.copy2(tmp_docx, output_path)

    size_kb = output_path.stat().st_size / 1024
    logging.info(f"Word document written to {output_path} ({size_kb:.0f} KB)")
    return output_path
