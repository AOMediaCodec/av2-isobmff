"""Main Word-to-Bikeshed converter orchestrating all extraction modules.

Pipeline:

1. Load the document with python-docx.
2. Extract metadata (title, author, date).
3. Walk the document body in element order (paragraphs and tables
   interleaved) and classify each via the style map.
4. Detect SDL syntax tables vs. data tables.
5. Extract figures and images.
6. Split content into sections at heading boundaries.
7. For each section generate ``.bs`` file content.
8. Generate ``header.bs``, ``manifest.txt``, ``symbols.bs``,
   ``bibliography.bs``, and a conversion report.
"""

from __future__ import annotations

import html as html_mod
import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING

from specbuild.input.equationextract import extract_equation, format_equation_bs
from specbuild.input.figureextract import (
    extract_images,
    format_figure_bs,
)
from specbuild.input.report import ConversionReport, generate_report_html, generate_report_text
from specbuild.input.sdldetect import detect_table_syntax, reconstruct_sdl, reconstruct_syntax_table
from specbuild.input.stylemap import (
    build_style_map,
    classify_with_map,
    detect_flavor,
    heading_level_in_map,
    is_heading_in_map,
    should_skip_in_map,
)
from specbuild.input.utils import W_TAG as _W_TAG
from specbuild.input.utils import make_html_id as _make_heading_id
from specbuild.input.utils import sanitize_filename as _sanitize_filename

#: Relationship namespace used on w:hyperlink r:id attributes.
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

if TYPE_CHECKING:
    from docx.document import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

#: Monospace font families used to detect inline code.
_MONO_FONTS = frozenset(
    {
        "Courier New",
        "Courier",
        "Consolas",
        "Menlo",
        "Monaco",
        "Source Code Pro",
        "Lucida Console",
        "DejaVu Sans Mono",
    }
)

#: Matches upper-case constant definitions: NAME = value
_CONST_DEF_RE = re.compile(r"^([A-Z][A-Z_0-9]+)\s*=\s*(\d+)")

#: Strips a Word-style caption number prefix, e.g. "Table 7-1 — ", "Figure 6.2 — ",
#: "Table 5 — ", "Figure C.1 — " (annex-letter.digit format).
#: renumber_annexes.py adds its own "Table X.N: " prefix at build time.
_CAPTION_NUMBER_RE = re.compile(r"^(?:Table|Figure)\s+(?:[A-Z]\.)?\d+(?:[-–.]\d+)?\s*[—–]\s*")

#: Extracts the kind ("table"/"figure") and number from a caption prefix.
#: Matches the same forms as _CAPTION_NUMBER_RE but captures the number so
#: an id like ``table-5-1`` or ``figure-c-1`` can be derived for cross-refs.
_CAPTION_ID_RE = re.compile(
    r"^(Table|Figure)\s+([A-I]?[\-\.]?\d+(?:[\-–\.]\d+)*)\b",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _is_informal_heading(paragraph) -> bool:
    """Return True if *paragraph* looks like an informally-styled heading.

    ISO publications sometimes use a bold "Body Text" or "Normal" paragraph
    as a visual sub-section heading without assigning it a proper Heading N
    style (e.g. "Publication and versions of this document" in the Introduction).
    This heuristic promotes those paragraphs to h3 headings.

    Criteria (all must hold):
    - Short text (≤ 120 chars — longer runs are likely sentences, not titles)
    - Every non-whitespace run is bold
    - No sentence-ending punctuation at the end (. ; : → likely a title, not prose)
    - Style is a generic body style (not already a code/note/dfn type)
    """
    text = paragraph.text.strip()
    if not text or len(text) > 120:
        return False
    # Must not end with period/semicolon/colon (would be prose, not title)
    if text[-1] in (".", ";", ":"):
        return False
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    # Every non-whitespace run must be bold
    return all(bool(r.bold) for r in runs)


def _strip_caption_number(text: str) -> str:
    return _CAPTION_NUMBER_RE.sub("", text)


def _caption_to_id(text: str) -> str:
    """Extract a normalized HTML id from a table/figure caption.

    Returns ``"table-5-1"``, ``"figure-c-1"``, etc., matching the IDs
    produced by xrefmap.py.  Returns ``""`` if no recognizable prefix.
    """
    m = _CAPTION_ID_RE.match(text)
    if not m:
        return ""
    kind = m.group(1).lower()
    # Normalize the number portion: en-dashes → ASCII hyphens, dots → hyphens,
    # lowercase letters, collapse to a slug.
    num = m.group(2).replace("–", "-").replace(".", "-").lower()
    return f"{kind}-{num}"


def _apply_run_formatting(text: str, run) -> str:
    """Apply bold/italic/mono/sub/sup formatting to an already-escaped text fragment.

    Uses HTML tags rather than markdown asterisks so the result renders
    correctly inside headings, list items, table cells, and other contexts
    where Bikeshed/markdown processing of inline `**`/`*` is unreliable.

    Leading and trailing whitespace is stripped from the text BEFORE wrapping
    in HTML tags and returned separately, so the formatted output never has
    whitespace *inside* the tag (``<em> NAL units</em>`` → ``<em>NAL units</em>``
    with the space restored outside the tag by the caller).

    Returns the fully-formatted fragment including surrounding whitespace.
    """
    # Preserve surrounding whitespace outside the tag
    lstrip = len(text) - len(text.lstrip())
    rstrip = len(text) - len(text.rstrip())
    leading = text[:lstrip]
    trailing = text[len(text) - rstrip :] if rstrip else ""
    core = text[lstrip : len(text) - rstrip if rstrip else None]

    if not core:
        # Space-only run with formatting — return as plain whitespace; the
        # formatting on whitespace alone has no semantic meaning.
        return text

    is_mono = run.font and run.font.name and run.font.name in _MONO_FONTS
    if is_mono:
        core = f"<code>{core}</code>"
    if run.bold:
        core = f"<strong>{core}</strong>"
    if run.italic:
        core = f"<em>{core}</em>"
    if run.font and run.font.subscript:
        core = f"<sub>{core}</sub>"
    if run.font and run.font.superscript:
        core = f"<sup>{core}</sup>"
    return leading + core + trailing


def _format_inline(paragraph: Paragraph) -> str:
    """Convert a paragraph's content into Bikeshed-flavored inline markup.

    Iterates the raw XML element tree so that ``w:hyperlink`` elements —
    invisible to ``paragraph.runs`` — are captured and their display text
    preserved.  External HTTP links become ``<a href="…">`` anchors; internal
    Word bookmarks are rendered as plain text (the display text Word computed
    already contains the resolved number, e.g. ``3.1.1``).

    Adjacent runs with identical formatting are coalesced so a phrase like
    ``"Publication and versions of " + "this document"`` (both bold) becomes
    one ``<strong>Publication and versions of this document</strong>`` span
    rather than two separate spans that would produce odd whitespace and
    redundant tags.
    """
    from docx.text.run import Run as _Run

    parts: list[str] = []
    # Buffer of adjacent runs with the same formatting signature
    buf_text = ""
    buf_run = None  # last run in the buffer (provides formatting)

    def _fmt_key(run) -> tuple:
        """Hashable formatting signature for a run."""
        font_name = run.font.name if run.font else None
        return (
            bool(run.bold),
            bool(run.italic),
            bool(run.font and run.font.subscript),
            bool(run.font and run.font.superscript),
            font_name in _MONO_FONTS if font_name else False,
        )

    def _flush_buf():
        nonlocal buf_text, buf_run
        if buf_text and buf_run is not None:
            parts.append(_apply_run_formatting(html_mod.escape(buf_text), buf_run))
        buf_text = ""
        buf_run = None

    for child in paragraph._element:
        local = child.tag.rpartition("}")[2]

        if local == "r":
            run = _Run(child, paragraph)
            text = run.text
            if not text:
                continue
            if buf_run is not None and _fmt_key(run) == _fmt_key(buf_run):
                buf_text += text
            else:
                _flush_buf()
                buf_text = text
                buf_run = run

        elif local == "hyperlink":
            _flush_buf()
            link_text = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
            if not link_text:
                continue
            escaped = html_mod.escape(link_text)
            r_id = child.get(f"{{{_R_NS}}}id", "")
            try:
                if r_id and r_id in paragraph.part.rels:
                    url = paragraph.part.rels[r_id].target_ref
                    if url.startswith("http"):
                        parts.append(f'<a href="{html_mod.escape(url)}">{escaped}</a>')
                    else:
                        parts.append(escaped)
                else:
                    parts.append(escaped)
            except (KeyError, AttributeError):
                parts.append(escaped)

    _flush_buf()
    return "".join(parts)


def _table_to_html(table: Table, caption: str = "", table_id: str = "") -> str:
    """Convert a python-docx data table to an HTML ``<table>`` string.

    Handles merged cells via the ``gridSpan`` attribute on ``<w:tcPr>``.
    Caption is embedded as ``<caption>`` inside the table when provided.
    When *table_id* is non-empty, it is emitted as the ``id`` attribute
    on the ``<table>`` element so cross-references to ``Table N-M`` resolve.

    Header row detection: rows marked with ``<w:tblHeader>`` in their
    ``<w:trPr>`` element are wrapped in ``<thead>`` with ``<th>`` cells.
    All subsequent rows go into ``<tbody>`` with ``<td>`` cells.
    Tables with no header rows emit only ``<tbody>``.
    """
    open_tag = (
        f'<table class="data" id="{html_mod.escape(table_id, quote=True)}">'
        if table_id
        else '<table class="data">'
    )
    lines: list[str] = [open_tag]

    if caption:
        lines.append(f"  <caption>{html_mod.escape(caption)}</caption>")

    def _row_is_header(row) -> bool:
        """Return True only if Word marked this row as a table header (tblHeader OPC flag).

        The bold-detection fallback is intentionally omitted: a bold first data
        cell (common in ISO tables) would otherwise produce false-positive headers.
        """
        try:
            tr_pr = row._tr.find(f"{_W_TAG}trPr")
        except AttributeError:
            return False
        if tr_pr is None:
            return False
        return tr_pr.find(f"{_W_TAG}tblHeader") is not None

    in_header = True
    section_opened = False

    for row in table.rows:
        is_hdr = _row_is_header(row)

        # Open the appropriate section element before the first row
        if not section_opened:
            if is_hdr:
                lines.append("  <thead>")
            else:
                lines.append("  <tbody>")
                in_header = False
            section_opened = True

        # Transition from header section to body section
        elif in_header and not is_hdr:
            lines.append("  </thead>")
            lines.append("  <tbody>")
            in_header = False

        lines.append("  <tr>")
        tag = "th" if is_hdr else "td"
        seen_tcs: set[int] = set()

        for cell in row.cells:
            # python-docx returns the same _tc element for every virtual cell
            # in a merged region — skip duplicates to avoid repeated output.
            tc_id = id(cell._tc)
            if tc_id in seen_tcs:
                continue
            seen_tcs.add(tc_id)

            tc = cell._tc
            grid_span_el = tc.find(f"{_W_TAG}tcPr/{_W_TAG}gridSpan")
            colspan = ""
            if grid_span_el is not None:
                span_val = grid_span_el.get(f"{_W_TAG}val", "1")
                if span_val and int(span_val) > 1:
                    colspan = f' colspan="{span_val}"'

            cell_text = html_mod.escape(cell.text.strip())
            lines.append(f"    <{tag}{colspan}>{cell_text}</{tag}>")

        lines.append("  </tr>")

    # Close open thead (table had only header rows) or tbody
    if section_opened:
        if in_header:
            lines.append("  </thead>")
        else:
            lines.append("  </tbody>")

    lines.append("</table>")
    lines.append("")
    return "\n".join(lines)


def _extract_metadata(doc: Document) -> dict[str, str]:
    """Extract document metadata (title, author, date) from core properties.

    ISO publications frequently ship with stale ``core.title`` values left
    over from prior ITU-T parallel-text edits (e.g. ``ITU-T Rec. H.265
    (04/2013)`` on a 2025 ISO/IEC 23008-2 FDIS).  When the cover page contains
    a clearly newer identifier in ``zzCover`` / ``zzSTDTitle`` paragraphs,
    prefer that.
    """
    props = doc.core_properties
    title = props.title or ""

    # Look for an ISO/IEC identifier and human-readable title on the cover page.
    iso_id: str | None = None
    iso_title: str | None = None
    for p in doc.paragraphs[:30]:
        if p.style.name not in ("zzCover", "zzSTDTitle", "zzCopyright"):
            continue
        text = p.text.strip()
        if not text:
            continue
        if iso_id is None and re.match(r"^(?:ISO/IEC|ISO|IEC)\s+[\d\-]+", text):
            iso_id = text
        elif iso_title is None and ("—" in text or "—" in text) and len(text) > 30:
            iso_title = text
        if iso_id and iso_title:
            break

    if iso_id and iso_title:
        # Combine identifier (without ":202X(en)" suffix) and descriptive title
        clean_id = re.sub(r":[\d\w()]+$", "", iso_id)
        title = f"{clean_id} - {iso_title}"
    elif iso_title:
        title = iso_title

    return {
        "title": title,
        "author": props.author or "",
        "created": str(props.created or ""),
        "modified": str(props.modified or ""),
        "subject": props.subject or "",
    }


def _generate_header_bs(
    metadata: dict[str, str],
    flavor: str,
    output_dir: Path,
) -> Path:
    """Generate the ``header.bs`` file with Bikeshed metadata."""
    title = metadata.get("title", "Imported Specification")
    author = metadata.get("author", "")

    lines: list[str] = [
        "<pre class='metadata'>",
        f"Title: {title}",
        "Status: WD",
        "Work Status: exploring",
        f"Shortname: {_make_heading_id(title) or 'imported-spec'}",
        "Level: 1",
        "URL: https://example.com/spec",
        f"Editor: {author}" if author else "Editor: TBD",
        "Abstract: This specification was imported from a Word document.",
        "Markup Shorthands: markdown yes",
        "</pre>",
        "",
        '<link rel="stylesheet" href="css/custom.css">',
        '<link rel="stylesheet" href="css/print.css" media="print">',
        "",
        "<script>",
        "MathJax = {",
        "  tex: {",
        "    inlineMath: [['\\\\(', '\\\\)']],",
        "    displayMath: [['$$', '$$']],",
        "  },",
        "  chtml: { scale: 1.0, mtextInheritFont: true }",
        "};",
        "</script>",
        '<script async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-chtml.js"></script>',
        "",
    ]

    if flavor:
        lines.insert(-1, f"<!-- Standards flavor: {flavor} -->")

    header_path = output_dir / "header.bs"
    header_path.write_text("\n".join(lines), encoding="utf-8")
    return header_path


def _generate_specbuild_toml(
    metadata: dict[str, str],
    flavor: str,
    docx_path: Path,
    output_dir: Path,
) -> Path:
    """Write a minimal specbuild.toml so the output can be built with compile.py."""
    title = metadata.get("title", "") or metadata.get("subject", "")

    # Extract ISO/ITU doc number from filename (e.g. "23000-19", "14496-12")
    doc_number = ""
    m = re.search(r"\b(\d{4,5}-\d+)\b", docx_path.stem)
    if m:
        doc_number = m.group(1)

    # Derive a short spec_name: last token of title, or stem fallback
    spec_name = ""
    if title:
        spec_name = title.split()[-1][:24].strip(".,;:")
    if not spec_name:
        # Strip edition/ID noise from filename stem
        stem = re.sub(r";.*$|\s+ed\b.*|\s+id\b.*", "", docx_path.stem).strip()
        spec_name = re.sub(r"[^A-Za-z0-9_-]", "_", stem)[:24]

    standards_flavor = {"cmaf": "iso", "h265": "itu-t"}.get(flavor, "iso")

    # Root-level SpecConfig fields must come BEFORE any [section] header in TOML
    toml_lines = [
        f'spec_name = "{spec_name}"',
        'bikeshed_dir = "."',
    ]
    if title:
        toml_lines.append(f'spec_full_name = "{title.replace(chr(34), chr(39))}"')
    toml_lines += [
        "",
        "[standards]",
        f'flavor = "{standards_flavor}"',
    ]
    if doc_number:
        toml_lines.append(f'docnumber = "{doc_number}"')
    toml_lines.append("")

    toml_path = output_dir / "specbuild.toml"
    toml_path.write_text("\n".join(toml_lines), encoding="utf-8")
    return toml_path


def _parse_annex_text(raw_text: str) -> tuple[str, str]:
    """Parse ANNEX-style paragraph text into (obligation, title).

    The raw text is typically ``"\\n(normative)\\n\\nTitle Text"``.
    """
    text = raw_text.strip()
    obligation = "normative"
    title = text

    m = re.match(r"\((\w+)\)\s*(.*)", text, re.DOTALL)
    if m:
        obligation = m.group(1).lower()
        title = m.group(2).strip()
    else:
        parts = text.split("\n")
        for part in parts:
            part = part.strip()
            if part.startswith("(") and part.endswith(")"):
                obligation = part[1:-1].lower()
            elif part and not part.startswith("("):
                title = part
                break

    return obligation, title


# ---------------------------------------------------------------------------
# Body element iteration
# ---------------------------------------------------------------------------


def _iter_body_elements(doc: Document):
    """Yield (element_type, element) tuples in document order.

    Walks ``doc.element.body`` children and classifies each as either
    a paragraph or a table, yielding the corresponding python-docx object.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body:
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, doc))


# ---------------------------------------------------------------------------
# Section splitting
# ---------------------------------------------------------------------------


def _split_into_sections(
    elements: list[tuple[str, object]],
    split_level: int,
    report: ConversionReport,
    style_map: dict | None = None,
) -> list[dict]:
    """Split interleaved elements into sections at heading boundaries.

    Args:
        elements:    List of (type_str, element) from :func:`_iter_body_elements`.
        split_level: Heading level at which to split (1 = Heading 1).
        report:      Conversion report to update.
        style_map:   Effective style map (from :func:`build_style_map`).

    Returns:
        List of section dicts with keys: ``heading_text``, ``heading_id``,
        ``heading_level``, ``elements`` (list of (type_str, element) pairs).
    """
    if style_map is None:
        from specbuild.input.stylemap import STYLE_MAP

        style_map = STYLE_MAP

    sections: list[dict] = []
    current: dict | None = None
    annex_counter = 0
    _annex_letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    for elem_type, elem in elements:
        if elem_type == "paragraph":
            style_name = getattr(getattr(elem, "style", None), "name", None)

            if should_skip_in_map(style_name, style_map):
                continue

            if is_heading_in_map(style_name, style_map):
                level = heading_level_in_map(style_name, style_map)
                if level <= split_level:
                    heading_text = elem.text.strip()

                    style_info = classify_with_map(style_name, style_map)
                    is_annex = style_info.get("is_annex", False)
                    annex_letter = ""
                    annex_obligation = ""
                    if is_annex:
                        annex_letter = (
                            _annex_letters[annex_counter]
                            if annex_counter < len(_annex_letters)
                            else str(annex_counter + 1)
                        )
                        annex_counter += 1
                        obligation, title = _parse_annex_text(heading_text)
                        annex_obligation = obligation
                        heading_text = f"Annex {annex_letter} ({obligation}) — {title}"

                    heading_id = _make_heading_id(heading_text)
                    current = {
                        "heading_text": heading_text,
                        "heading_id": heading_id,
                        "heading_level": level,
                        "is_annex": is_annex,
                        "annex_letter": annex_letter,
                        "annex_obligation": annex_obligation,
                        "annex_title": title if is_annex else "",
                        "elements": [],
                    }
                    sections.append(current)
                    report.sections_generated += 1
                    continue

        # Add to current section (or create a preamble section)
        if current is None:
            current = {
                "heading_text": "Preamble",
                "heading_id": "preamble",
                "heading_level": 0,
                "is_annex": False,
                "annex_letter": "",
                "annex_obligation": "",
                "elements": [],
            }
            sections.append(current)

        current["elements"].append((elem_type, elem))

    return sections


# ---------------------------------------------------------------------------
# Section rendering
# ---------------------------------------------------------------------------


def _render_section(
    section: dict,
    images: dict[str, Path],
    detect_sdl: bool,
    report: ConversionReport,
    syntax_format: str = "table",
    style_map: dict | None = None,
    flavor: str = "",
    custom_descriptor_re: re.Pattern | None = None,
) -> str:
    """Render a section's elements to Bikeshed source text.

    Args:
        section:             Section dict from :func:`_split_into_sections`.
        images:              Image relationship map from :func:`extract_images`.
        detect_sdl:          Whether to auto-detect SDL tables.
        report:              Conversion report to update.
        syntax_format:       ``"table"`` or ``"sdl"`` code block output.
        style_map:           Effective style map (from :func:`build_style_map`).
        flavor:              Standards flavor for SDL dispatch.
        custom_descriptor_re: Optional override for descriptor detection regex.

    Returns:
        Bikeshed ``.bs`` source text for this section.
    """
    if style_map is None:
        from specbuild.input.stylemap import STYLE_MAP

        style_map = STYLE_MAP
    lines: list[str] = []
    heading = section["heading_text"]
    level = section["heading_level"]

    if level > 0:
        hashes = "#" * min(level, 6)
        lines.append(f"{hashes} {heading} {hashes} {{#{section['heading_id']}}}")
        lines.append("")

    in_note = False
    in_list = False
    list_type = ""
    in_dfn_list = False
    in_ref_list = False
    # Buffer for TermNum paragraph — held until the following Term(s) paragraph
    # arrives so they can be merged into one <dt> element.
    pending_dfn_num = ""
    pending_table_caption = ""
    pending_table_id = ""
    pending_figure_caption = ""
    pending_figure_id = ""
    pending_figure: dict | None = (
        None  # deferred until caption arrives (ISO: caption after graphic)
    )
    code_buffer: list[str] = []

    def _flush_code() -> None:
        if code_buffer:
            lines.append("```")
            lines.extend(code_buffer)
            lines.append("```")
            lines.append("")
            code_buffer.clear()

    def _close_dfn_list() -> None:
        nonlocal in_dfn_list, pending_dfn_num
        if pending_dfn_num:
            lines.append(f"<dt>{pending_dfn_num}</dt>")
            pending_dfn_num = ""
        if in_dfn_list:
            lines.append("</dl>")
            lines.append("")
            in_dfn_list = False

    def _close_ref_list() -> None:
        nonlocal in_ref_list
        if in_ref_list:
            lines.append("</ul>")
            lines.append("")
            in_ref_list = False

    def _flush_pending_figure(caption: str = "") -> None:
        nonlocal pending_figure
        if pending_figure is None:
            return
        pending_figure["caption"] = caption
        lines.append(format_figure_bs(pending_figure))
        report.total_figures += 1
        pending_figure = None

    def _flush_orphan_caption() -> None:
        """Emit a caption-only ``<figure>`` for an orphan caption.

        When a ``Figure title`` paragraph is processed but no ``Figure Graphic``
        ever follows (common in ISO "text-clean" editions where figures are
        stripped), we still need an id-anchored placeholder so cross-references
        like ``[Figure 6-1]`` resolve.
        """
        nonlocal pending_figure_caption, pending_figure_id
        if pending_figure_caption and pending_figure_id:
            lines.append(f'<figure id="{pending_figure_id}">')
            lines.append(f"  <figcaption>{html_mod.escape(pending_figure_caption)}</figcaption>")
            lines.append("</figure>")
            lines.append("")
            report.total_figures += 1
        pending_figure_caption = ""
        pending_figure_id = ""

    for elem_type, elem in section["elements"]:
        if elem_type == "paragraph":
            style_name = getattr(getattr(elem, "style", None), "name", None)
            info = classify_with_map(style_name, style_map)
            para_type = info.get("type", "paragraph")
            text = elem.text.strip()

            report.total_paragraphs += 1

            # Flush buffered code when leaving code style
            if para_type != "code":
                _flush_code()

            # Close note block if we leave note style
            if in_note and para_type != "note":
                lines.append("</div>")
                lines.append("")
                in_note = False

            # Close list if style changes
            if in_list and para_type != "list_item":
                close_tag = "</ol>" if list_type == "ordered" else "</ul>"
                lines.append(close_tag)
                lines.append("")
                in_list = False

            # Close definition list if leaving dfn-related styles
            if in_dfn_list and para_type not in ("dfn", "dfn_num", "definition"):
                _close_dfn_list()

            # Close reference list if leaving ref/biblio styles
            if in_ref_list and para_type not in ("biblio", "ref_norm"):
                _close_ref_list()

            if para_type == "heading":
                _flush_orphan_caption()
                h_level = min(info.get("level", 2), 6)
                h_id = _make_heading_id(text)
                hashes = "#" * h_level
                lines.append("")
                lines.append(f"{hashes} {text} {hashes} {{#{h_id}}}")
                lines.append("")

            elif para_type == "equation":
                eq = extract_equation(elem)
                bs_eq = format_equation_bs(eq)
                if bs_eq:
                    lines.append(bs_eq)
                    # Only warn about OMML/embedded extraction when we DIDN'T
                    # successfully pull a rendered preview image — when we have
                    # the image, the equation displays correctly.
                    if (eq.get("has_omml") or eq.get("has_image")) and not eq.get("image_path"):
                        lines.append("")
                        lines.append(
                            '<p class="issue">'
                            "EDITOR: This equation was extracted from an embedded object "
                            "and may need manual formatting. Verify the mathematical "
                            "notation is correct.</p>"
                        )
                elif text:
                    lines.append(f"<p>{text}</p>")
                    lines.append("")
                    lines.append(
                        '<p class="issue">'
                        "EDITOR: This equation could not be fully extracted. "
                        "Please add the correct mathematical notation.</p>"
                    )
                report.equations_extracted += 1

            elif para_type == "code":
                code_buffer.append(text)

            elif para_type == "example":
                inline = _format_inline(elem)
                lines.append('<div class="example">')
                lines.append(inline)
                lines.append("</div>")
                lines.append("")

            elif para_type == "dfn":
                if not in_dfn_list:
                    lines.append("<dl>")
                    in_dfn_list = True
                inline = _format_inline(elem)
                # Merge with preceding TermNum (e.g. "3.7") into one <dt> so
                # the term number and term name are not split across two <dt>s.
                if pending_dfn_num:
                    lines.append(f"<dt>{pending_dfn_num} <dfn>{inline}</dfn></dt>")
                    pending_dfn_num = ""
                else:
                    lines.append(f"<dt><dfn>{inline}</dfn></dt>")
                report.terms_extracted += 1

            elif para_type == "dfn_num":
                if not in_dfn_list:
                    lines.append("<dl>")
                    in_dfn_list = True
                # Buffer until the Term(s) paragraph follows so we can combine
                # them.  If the next paragraph is NOT a dfn, flush standalone.
                inline = _format_inline(elem)
                if pending_dfn_num:
                    # Two consecutive dfn_num without dfn — emit the previous one
                    lines.append(f"<dt>{pending_dfn_num}</dt>")
                pending_dfn_num = inline

            elif para_type == "definition":
                if pending_dfn_num:
                    # dfn_num was never followed by a dfn — emit it standalone
                    lines.append(f"<dt>{pending_dfn_num}</dt>")
                    pending_dfn_num = ""
                inline = _format_inline(elem)
                if in_dfn_list:
                    lines.append(f"<dd>{inline}</dd>")
                else:
                    lines.append(f"<p>{inline}</p>")
                    lines.append("")

            elif para_type == "note":
                if not in_note:
                    lines.append('<div class="note">')
                    in_note = True
                inline = _format_inline(elem)
                lines.append(inline)

            elif para_type == "list_item":
                lt = info.get("list_type", "bullet")
                if not in_list or list_type != lt:
                    # Close previous list if type changed
                    if in_list:
                        close_tag = "</ol>" if list_type == "ordered" else "</ul>"
                        lines.append(close_tag)
                    open_tag = "<ol>" if lt == "ordered" else "<ul>"
                    lines.append(open_tag)
                    in_list = True
                    list_type = lt
                inline = _format_inline(elem)
                # Strip leading em-dash/en-dash bullet character Word puts in text
                inline = re.sub(r"^[—–]\s*", "", inline)
                lines.append(f"  <li>{inline}</li>")

            elif para_type == "table_caption":
                _flush_pending_figure()  # caption didn't follow — emit without one
                _flush_orphan_caption()  # flush any orphan figure caption first
                pending_table_caption = _strip_caption_number(text)
                pending_table_id = _caption_to_id(text)
                pending_figure_caption = ""
                pending_figure_id = ""

            elif para_type == "figure_caption":
                caption_text = _strip_caption_number(text)
                caption_id = _caption_to_id(text)
                if pending_figure is not None:
                    # Caption arrived after its graphic — emit both together
                    if caption_id:
                        pending_figure["figure_id"] = caption_id
                    _flush_pending_figure(caption_text)
                else:
                    # If a previous caption was orphaned (no graphic followed it),
                    # flush it as a caption-only figure so its ID is anchored,
                    # then store the new caption.
                    if pending_figure_caption and caption_id != pending_figure_id:
                        _flush_orphan_caption()
                    pending_figure_caption = caption_text
                    pending_figure_id = caption_id
                pending_table_caption = ""

            elif para_type == "figure_graphic":
                from specbuild.input.figureextract import _extract_image_rid

                _flush_pending_figure()  # flush any prior un-captioned figure
                rid = _extract_image_rid(elem)
                img_path = images.get(rid) if rid else None
                if img_path:
                    # Default sequential ID when no caption is found; will be
                    # overridden by _caption_to_id when the caption arrives.
                    default_fig_id = f"figure-{report.total_figures + 1}"
                    if pending_figure_caption:
                        # Caption already seen (before graphic) — use its derived ID
                        pending_figure = {
                            "image_path": img_path,
                            "caption": pending_figure_caption,
                            "figure_id": pending_figure_id or default_fig_id,
                        }
                        pending_figure_caption = ""
                        pending_figure_id = ""
                        _flush_pending_figure()
                    else:
                        # Defer until next figure_caption paragraph
                        pending_figure = {
                            "image_path": img_path,
                            "caption": "",
                            "figure_id": default_fig_id,
                        }

            elif para_type == "biblio":
                if not in_ref_list:
                    lines.append("<ul>")
                    in_ref_list = True
                inline = _format_inline(elem)
                lines.append(f"  <li>{inline}</li>")
                report.bibliography_entries += 1

            elif para_type == "ref_norm":
                if not in_ref_list:
                    lines.append("<ul>")
                    in_ref_list = True
                inline = _format_inline(elem)
                lines.append(f"  <li>{inline}</li>")
                report.bibliography_entries += 1

            elif para_type == "skip":
                pass

            else:
                # Default: paragraph.
                # Heuristic: a short paragraph where ALL text runs are bold and
                # the style is a generic body style acts as an informal heading
                # (common in ISO publications for introduction sub-sections like
                # "Publication and versions of this document").  Promote it to
                # ## (h2 in Bikeshed output) — one level below the top-level
                # # sections so the TOC hierarchy stays valid.
                if text and _is_informal_heading(elem):
                    h_id = _make_heading_id(text)
                    lines.append("")
                    lines.append(f"## {text} ## {{#{h_id}}}")
                    lines.append("")
                elif text:
                    inline = _format_inline(elem)
                    lines.append(inline)
                    lines.append("")

        elif elem_type == "table":
            # Flush any buffered code or pending figure before a table
            _flush_code()
            _close_dfn_list()
            _close_ref_list()
            _flush_pending_figure()
            report.total_tables += 1

            if detect_sdl:
                sdl_kind = detect_table_syntax(elem, flavor or "auto", custom_descriptor_re)
            else:
                sdl_kind = None

            if sdl_kind == "video_syntax":
                if syntax_format == "sdl":
                    sdl_block = reconstruct_sdl(elem)
                    lines.append(sdl_block)
                else:
                    syntax_html = reconstruct_syntax_table(elem)
                    lines.append(syntax_html)
                lines.append("")
                report.sdl_tables_detected += 1
            else:
                # Data table — embed caption inside table, prefer table_caption
                caption_text = pending_table_caption or pending_figure_caption
                caption_id = pending_table_id or pending_figure_id
                pending_table_caption = ""
                pending_table_id = ""
                pending_figure_caption = ""
                pending_figure_id = ""
                lines.append(_table_to_html(elem, caption_text, table_id=caption_id))

    # Close any open blocks
    _flush_code()
    _flush_pending_figure()
    _flush_orphan_caption()
    _close_dfn_list()
    _close_ref_list()
    if in_note:
        lines.append("</div>")
        lines.append("")
    if in_list:
        close_tag = "</ol>" if list_type == "ordered" else "</ul>"
        lines.append(close_tag)
        lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Symbol extraction
# ---------------------------------------------------------------------------


def _extract_symbols(doc: Document) -> dict[str, str]:
    """Scan paragraphs for constant definitions (NAME = value)."""
    symbols: dict[str, str] = {}
    for para in doc.paragraphs:
        m = _CONST_DEF_RE.match(para.text.strip())
        if m:
            symbols[m.group(1)] = m.group(2)
    return symbols


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------


def convert_docx(
    docx_path: Path,
    output_dir: Path,
    *,
    split_level: int = 1,
    detect_sdl: bool = True,
    extract_symbols: bool = True,
    extract_images_flag: bool = True,
    flavor: str = "",
    syntax_format: str = "table",
    custom_descriptor_re: re.Pattern | None = None,
) -> dict:
    """Convert a Word document to Bikeshed source files.

    Args:
        docx_path:            Path to the ``.docx`` file.
        output_dir:           Directory to write generated ``.bs`` files.
        split_level:          Split at Heading N level (default: 1).
        detect_sdl:           Auto-detect SDL syntax tables.
        extract_symbols:      Extract constants for ``symbols.bs``.
        extract_images_flag:  Extract embedded images.
        flavor:               Standards flavor hint (``"h265"``, ``"cmaf"``, or
                              ``""``/``"auto"`` to auto-detect).
        syntax_format:        ``"table"`` (HTML table) or ``"sdl"`` (fenced block).
        custom_descriptor_re: Override the SDL descriptor detection regex.

    Returns:
        Dict with keys: ``bs_files``, ``manifest_path``, ``report``.
    """
    try:
        from docx import Document
    except ImportError:
        logging.error(
            "python-docx is required for Word import. Install it: pip install python-docx"
        )
        raise SystemExit(1)

    logging.info(f"Importing Word document: {docx_path}")
    report = ConversionReport()

    # 1. Load document
    doc = Document(str(docx_path))

    # 2. Auto-detect flavor if not provided
    if not flavor or flavor == "auto":
        flavor = detect_flavor(doc)
        logging.info(f"Auto-detected document flavor: {flavor}")
    else:
        logging.info(f"Using document flavor: {flavor}")

    style_map = build_style_map(flavor)

    # 3. Extract metadata
    metadata = _extract_metadata(doc)
    logging.info(f"Document title: {metadata.get('title', '(none)')}")

    # 3. Extract images
    images: dict[str, Path] = {}
    if extract_images_flag:
        images = extract_images(doc, output_dir)
        report.images_extracted = len(images)

    # 4. Walk body elements in document order
    elements = list(_iter_body_elements(doc))
    logging.info(f"Document has {len(elements)} body elements")

    # 5. Split into sections
    sections = _split_into_sections(elements, split_level, report, style_map)
    logging.info(f"Split into {len(sections)} sections")

    # 6. Render each section to .bs content
    from specbuild.input.xrefmap import build_content_id_map, resolve_text_references

    output_dir.mkdir(parents=True, exist_ok=True)
    bs_files: list[Path] = []
    manifest_entries: list[str] = []

    # Generate header.bs
    header_path = _generate_header_bs(metadata, flavor, output_dir)
    bs_files.append(header_path)
    manifest_entries.append("header.bs")
    report.add_bs_file("header.bs")

    content_id_map = build_content_id_map(sections)

    # First pass: render all sections to collect figure IDs before resolving refs
    rendered: list[tuple[int, dict, str]] = []
    for idx, section in enumerate(sections):
        content = _render_section(
            section,
            images,
            detect_sdl,
            report,
            syntax_format,
            style_map=style_map,
            flavor=flavor,
            custom_descriptor_re=custom_descriptor_re,
        )
        rendered.append((idx, section, content))

    # Populate figure IDs in the cross-reference map from rendered content
    _fig_id_re = re.compile(r'<figure\s+id="fig-(\d+)"')
    for _idx, _section, content in rendered:
        for m in _fig_id_re.finditer(content):
            fig_num = m.group(1)
            content_id_map[f"figure-{fig_num}"] = f"fig-{fig_num}"

    # Second pass: resolve cross-references and write files
    total_resolved = 0
    for idx, section, content in rendered:
        heading = section["heading_text"]
        is_annex = section.get("is_annex", False)
        annex_letter = section.get("annex_letter", "")

        if is_annex and annex_letter:
            annex_title = section.get("annex_title", "")
            clean_title = _sanitize_filename(annex_title) if annex_title else "untitled"
            filename = f"{idx:03d}_annex_{annex_letter.lower()}_{clean_title}.bs"
        else:
            filename = f"{idx:03d}_{_sanitize_filename(heading)}.bs"

        content, n = resolve_text_references(content, content_id_map)
        total_resolved += n

        file_path = output_dir / filename
        file_path.write_text(content, encoding="utf-8")
        bs_files.append(file_path)
        manifest_entries.append(filename)
        report.add_bs_file(filename)

    report.xrefs_resolved = total_resolved
    if total_resolved:
        logging.info(f"Resolved {total_resolved} text-based cross-references")

    # 7. Extract symbols if requested
    if extract_symbols:
        symbols = _extract_symbols(doc)
        if symbols:
            symbols_lines = [
                "<!-- Extracted symbol constants -->",
                "<pre class='metadata'>",
                "<!-- Symbol definitions extracted from source document -->",
                "</pre>",
                "",
            ]
            for name, value in sorted(symbols.items()):
                symbols_lines.append(f"<pre>  {name} = {value}</pre>")
            symbols_lines.append("")

            symbols_path = output_dir / "symbols.bs"
            symbols_path.write_text("\n".join(symbols_lines), encoding="utf-8")
            bs_files.append(symbols_path)
            manifest_entries.append("symbols.bs")
            report.add_bs_file("symbols.bs")

    # 8. Generate manifest.txt with annex separation
    manifest_path = output_dir / "manifest.txt"
    main_entries = []
    annex_entries = []
    for entry in manifest_entries:
        if "_annex_" in entry:
            annex_entries.append(entry)
        else:
            main_entries.append(entry)

    manifest_lines = main_entries[:]
    if annex_entries:
        manifest_lines.append("")
        manifest_lines.append("# --- Annexes ---")
        manifest_lines.extend(annex_entries)

    manifest_path.write_text("\n".join(manifest_lines) + "\n", encoding="utf-8")
    logging.info(f"Wrote manifest.txt with {len(manifest_entries)} entries")

    # 9. Generate specbuild.toml
    toml_path = _generate_specbuild_toml(metadata, flavor, docx_path, output_dir)
    logging.info(f"Wrote specbuild.toml: {toml_path}")

    # 10. Generate conversion report
    report_text = generate_report_text(report)
    logging.info("\n" + report_text)

    report_html_path = output_dir / "conversion_report.html"
    generate_report_html(report, report_html_path)

    return {
        "bs_files": [str(p) for p in bs_files],
        "manifest_path": str(manifest_path),
        "toml_path": str(toml_path),
        "report": report,
    }
